from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "人工智能导论实验报告_任务2任务3_统一字体.docx"

doc = Document(str(DOCX))

for section in doc.sections:
    for footer in (section.footer, section.first_page_footer, section.even_page_footer):
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""

doc.save(str(DOCX))
print(DOCX)
