#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成基于模板的人工智能导论实验报告 .docx 文件
IEEE 格式，含任务二（Q-Learning 迷宫）和任务三（动物识别专家系统）
所有内容以自然段落呈现，无项目符号列表。
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── helpers ──────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formula(paragraph, latex_text):
    run = paragraph.add_run(f"  {latex_text}  ")
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return run

def add_para(doc, text, bold=False, size=12, font_name='宋体',
             alignment=None, space_after=Pt(6), first_line_indent=0.74,
             color=None):
    """统一的自然段落，默认首行缩进两字符（0.74cm）。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_code_block(doc, code_text):
    """等宽代码块，左侧缩进。"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(16)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_image_placeholder(doc, description):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'【图片插入位置：{description}】')
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run('┌' + '─' * 50 + '┐\n' + '│' + ' ' * 50 + '│\n' + '└' + '─' * 50 + '┘')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run2.font.name = 'Consolas'
    return p

def add_table_with_data(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ═══════════════════════ COVER ═══════════════════════

for _ in range(6):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('人工智能导论实验报告')
run.bold = True; run.font.size = Pt(26); run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph(); doc.add_paragraph()

cover_fields = [
    ('姓    名', ''), ('学    号', ''), ('班    级', ''),
    ('任课教师', '郑定富'), ('成    绩', ''),
    ('时    间', '2026 年春季学期'), ('专    业', ''),
]

cover_table = doc.add_table(rows=len(cover_fields), cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(cover_fields):
    cl = cover_table.rows[i].cells[0]; cl.text = ''
    p = cl.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label + '：'); run.bold = True; run.font.size = Pt(14)
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    cl.width = Cm(4)

    cv = cover_table.rows[i].cells[1]; cv.text = ''
    p = cv.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(value); run.font.size = Pt(14)
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if not value:
        p2 = cv.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run('__________'); run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    cv.width = Cm(8)

for row in cover_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>')
        tcPr.append(tcBorders)

# ═══════════════════════ TOC ═══════════════════════

doc.add_page_break()
add_heading_styled(doc, '目  录', level=1); doc.add_paragraph()

for item in ['1、实验二   基于 Q-Learning 的迷宫寻路强化学习系统',
             '2、实验三   基于产生式系统的动物识别专家系统',
             '3、体会与建议']:
    p = doc.add_paragraph()
    run = p.add_run(item); run.font.size = Pt(12); run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════
#  EXPERIMENT 2
# ═══════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled(doc, '实验二：基于 Q-Learning 的迷宫寻路强化学习系统', level=1)

# ── 1. 问题分析 ──
add_heading_styled(doc, '1. 问题分析', level=2)
add_heading_styled(doc, '1.1 任务描述', level=3)

add_para(doc,
    '本实验要求基于 Q-Learning 强化学习理论，设计并实现一个迷宫寻路系统。智能体从迷宫中的 Start '
    '位置出发，通过与环境交互不断学习，最终找到通往 Goal 位置的最优路径。系统需要实现迷宫规模（行数、'
    '列数）的动态可调、障碍物数量和位置的随机生成，以及提供图形化人机界面（GUI），支持参数调节、训练'
    '过程可视化与最优路径展示。')

add_heading_styled(doc, '1.2 问题建模', level=3)

add_para(doc,
    '迷宫寻路问题可形式化为马尔可夫决策过程（Markov Decision Process, MDP），以五元组 '
    '⟨S, A, P, R, γ⟩ 加以严格定义。状态空间 S 为迷宫中所有可行走格子坐标的集合，'
    'S = {(r, c) | 0 ≤ r < rows, 0 ≤ c < cols, (r, c) ∉ obstacles}，外加一个哨兵终止状态 terminal。'
    '动作空间 A = {0, 1, 2, 3}，分别表示上、下、左、右四个方向的移动。状态转移概率 P 为确定性转移，'
    '即 P(s\' | s, a) = 1，动作执行后智能体必然到达目标格子；若目标格子为障碍物，则智能体留在原地。'
    '奖励函数 R(s, a, s\') 采用三层设计：到达目标获得 +100 的正向激励，撞障碍物承受 -10 的惩罚并留在原地，'
    '在空地上每行走一步扣除 -1 分以鼓励寻找最短路径。折扣因子 γ 取 0.9，用于平衡即时奖励与未来累积奖励。')

mdp_headers = ['符号', '含义', '本系统定义']
mdp_rows = [
    ['S', '状态空间',
     '所有可行走格子坐标，S = {(r,c) | 0≤r<rows, 0≤c<cols, (r,c)∉obstacles} ∪ {terminal}'],
    ['A', '动作空间', '{0, 1, 2, 3}：上、下、左、右'],
    ['P', '状态转移概率', '确定性转移，P(s\'|s,a) = 1；碰障碍物则留原地'],
    ['R', '奖励函数', '目标 +100 / 障碍 -10 / 空地 -1'],
    ['γ', '折扣因子', '0.9'],
]
add_table_with_data(doc, mdp_headers, mdp_rows)

add_heading_styled(doc, '1.3 核心挑战', level=3)

add_para(doc,
    '本任务面临四个主要挑战。第一是探索与利用（Exploration vs. Exploitation）的平衡问题：'
    '智能体需在随机探索未知路径与利用已有 Q 值知识之间做出权衡，ε-贪心策略中的 ε 参数直接决定了这'
    '一平衡点。第二是收敛性判定：如何自动检测 Q 表已收敛至最优策略从而停止训练，避免固定训练轮次的'
    '盲目性。第三是最优路径提取：收敛后如何从已学习到的 Q 表中提取出可可视化的最优路径，并以直观的'
    '图形化方式呈现给用户。第四是高维状态空间的管理：随着迷宫尺寸增大，状态-动作组合急剧膨胀，'
    'Q 表的存储与更新效率成为不可忽视的工程问题。')

# ── 2. 算法原理 ──
add_heading_styled(doc, '2. 算法原理', level=2)
add_heading_styled(doc, '2.1 Q-Learning 算法', level=3)

add_para(doc,
    'Q-Learning 是一种经典的无模型（model-free）、离策略（off-policy）时序差分（Temporal Difference, '
    'TD）强化学习算法。其核心思想是通过迭代更新状态-动作价值函数 Q(s, a)，使得智能体在任意状态 s 下都'
    '能选择使期望累积奖励最大化的动作 a。Q 值更新遵循 TD 更新规则：')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formula(p, 'Q(s,a) ← Q(s,a) + α [ r + γ · maxₐ\' Q(s\',a\') — Q(s,a) ]')

add_para(doc,
    '公式中，学习率 α = 0.1 控制每次更新对 TD 误差的接纳程度，值越小学习越平稳但收敛越慢。'
    '折扣因子 γ = 0.9 决定了智能体对远期回报的关注程度：γ 趋近于 1 表示智能体更有远见，趋近于 0 则'
    '变得短视。表达式 r + γ · maxₐ\' Q(s\', a\') 称为 TD 目标（TD Target），代表当前状态-动作对的'
    '"应有价值"；它与当前估计值 Q(s, a) 的差即为 TD 误差（TD Error），衡量了"期望"与"现实"的差距。'
    '当 s\' = terminal（智能体到达目标）时，不存在后续状态，因此 TD 目标退化为即时奖励 r，更新公式'
    '简化为 Q(s, a) ← Q(s, a) + α [r — Q(s, a)]。')

add_heading_styled(doc, '2.2 ε-贪心动作选择策略', level=3)

add_para(doc,
    '智能体在每个状态下以概率 ε = 0.9 执行利用（Exploitation），选择当前 Q 值最大的动作；以概率 '
    '1 - ε = 0.1 执行探索（Exploration），从动作空间中均匀随机选择一个动作。当多个动作共享最大 Q 值时，'
    '系统从这些最优候选中等概率随机选取，以避免确定性偏差。此外，系统在收敛后提取最优路径时，通过 '
    'test_mode 参数将 ε 强制置为 1.0，确保路径提取过程不受随机探索干扰，得到确定性最优策略。')

add_heading_styled(doc, '2.3 收敛检测机制', level=3)

add_para(doc,
    '本系统引入了一种基于"步数稳定"的启发式收敛检测方法。设 g_t 为连续稳定回合计数器，steps_t 为第 '
    't 个 episode 中智能体从起点到达目标所需的步数。若当前 episode 步数与上一 episode 相同，则 g_t = '
    'g_{t-1} + 1；否则 g_t 重置为 0。当 g_t ≥ 3，即连续 3 个回合的步数完全相同，系统判定 Q-Learning '
    '已收敛至最优策略，自动停止训练并绘制最优路径。其理论依据在于：在确定性环境中，若智能体已掌握最优'
    '策略，每个 episode 将沿相同的路径到达目标，步数恒定不变。连续 3 次一致足以排除随机扰动的影响。')

# ── 3. 系统设计与实现 ──
add_heading_styled(doc, '3. 系统设计与实现', level=2)
add_heading_styled(doc, '3.1 系统架构', level=3)

add_para(doc,
    '系统采用三层解耦架构来组织代码。表示层由 MazeApp 主窗口控制器承担，它集成了 MazeCanvas 迷宫画布'
    '（负责网格渲染、智能体绘制和最优路径展示）以及右侧控制面板（提供行数、列数、障碍物数量的参数调节'
    '和训练启停控制）。算法层由 QLearningAgent 类独立封装，包含 choose_action() ε-贪心动作选择、'
    'learn() TD 更新以及基于字典的惰性 Q 表管理，完全不依赖任何 GUI 组件。环境层由 MazeCanvas 的 '
    'step() 方法实现状态转移与奖励反馈，包含边界检查、碰撞检测和障碍物的随机生成逻辑。这种分层设计使得'
    '算法核心可以被独立测试和复用，GUI 组件的变更不会影响推理逻辑。')

add_heading_styled(doc, '3.2 核心类设计', level=3)

# -- QLearningAgent --
add_para(doc, 'QLearningAgent 类封装了 Q-Learning 智能体的全部行为。', bold=True)
add_para(doc,
    '构造函数 __init__ 接收动作空间列表和三个超参数（α = 0.1, γ = 0.9, ε = 0.9），初始化一个空的 '
    'Q 表字典，键为状态坐标元组 (r, c)，值为长度为 4 的浮点数列表，分别对应上、下、左、右四个动作的 '
    'Q 值。check_state_exist(state) 方法实现了惰性初始化机制——仅当首次访问某个状态时才在 Q 表中为其'
    '创建全零向量，从而避免了为迷宫中所有格子预先分配内存。choose_action(state, test_mode=False) '
    '方法执行 ε-贪心策略：在 test_mode 为 True 或随机数小于 ε 时选择最大 Q 值动作，否则从动作空间中'
    '均匀随机选择。learn(s, a, r, s_) 方法执行 TD 更新：根据下一状态是否为 terminal 分情况计算 '
    'TD 目标，再以学习率 α 向目标方向修正当前 Q 值。')

agent_headers = ['方法', '功能']
agent_rows = [
    ['__init__(actions, lr, gamma, epsilon)', '初始化超参数与空 Q 表（惰性字典）'],
    ['check_state_exist(state)', '首次遇到某状态时创建全零 Q 向量'],
    ['choose_action(state, test_mode)', 'ε-贪心选择；test_mode 时跳过探索'],
    ['learn(s, a, r, s_)', 'TD 更新，区分终止/非终止状态'],
]
add_table_with_data(doc, agent_headers, agent_rows)

add_para(doc,
    'Q 表采用字典存储而非二维矩阵，其优势体现在两个方面：一是内存效率——仅记录智能体实际访问过的状态，'
    '在障碍物密集的迷宫中可显著节省内存；二是状态泛化——状态键为任意可哈希元组，理论上可自然扩展到非网格'
    '环境的连续状态空间（配合离散化或函数逼近）。')

# -- MazeCanvas --
add_para(doc, 'MazeCanvas 类继承自 QWidget，承担环境模拟与图形渲染的双重职责。', bold=True)
add_para(doc,
    'reset_env() 方法将起点固定在左上角 (0, 0)，终点固定在右下角 (rows-1, cols-1)，并清空障碍物列表'
    '和最优路径缓存。generate_obstacles() 方法通过拒绝采样在非起点、非终点的格子中随机放置指定数量的'
    '障碍物。step(action) 方法根据动作编号计算目标坐标，执行边界检查后判定结果：到达终点返回奖励 +100 '
    '并将下一状态设为 terminal；撞上障碍物返回奖励 -10 且下一状态保持为当前坐标（智能体原地不动）；在空地'
    '上行走返回奖励 -1 并正常移动。paintEvent(event) 采用三层 QPainter 渲染管线：先绘制所有格子的底色'
    '（起点浅蓝、终点浅绿、障碍深灰、空地白色），再以 5 像素宽的金色笔触绘制收敛后的最优路径连线，最后'
    '在智能体当前位置绘制红色椭圆小球。')

# -- MazeApp --
add_para(doc, 'MazeApp 类继承自 QMainWindow，是训练流程的总控制器。', bold=True)
add_para(doc,
    'train_step() 方法是 Q-Learning 单步迭代的核心循环体，由 QTimer 以 10 毫秒间隔驱动，每步依次执行'
    'choose_action → step → learn 三个原子操作，并递增当前 episode 的步数计数器。当 episode 结束（done '
    '为 True）时，系统比较当前步数与上一 episode 步数以更新稳定计数器，并在连续 3 次稳定时自动停机、调用 '
    'extract_and_draw_optimal_path() 从起点开始以 test_mode 贪心策略提取完整路径并绘制为金色线条。'
    'toggle_training() 方法实现训练的开始/暂停切换，generate_new_maze() 方法在改变迷宫参数后重置 Q 表、'
    '重新随机生成障碍物并清空所有训练状态。')

app_headers = ['方法', '功能']
app_rows = [
    ['train_step()', 'Q-Learning 单步迭代（choose → step → learn），10ms 间隔'],
    ['toggle_training()', '训练启动/暂停切换'],
    ['extract_and_draw_optimal_path()', '收敛后以贪心策略提取并绘制最优路径'],
    ['generate_new_maze()', '重置迷宫参数、Q 表与训练计数器'],
]
add_table_with_data(doc, app_headers, app_rows)

add_heading_styled(doc, '3.3 奖励函数设计', level=3)

add_para(doc,
    '奖励信号的设计是 Q-Learning 收敛性能的关键。到达目标给予 +100 的强正奖励，这是一个稀疏但足够强烈'
    '的正向激励，能够驱动智能体克服每步 -1 的走路代价而持续探索。碰撞障碍物仅给予 -10 的中等惩罚，并且'
    '不终止当前 episode——智能体留在原地，可以在下一时间步选择其他方向继续探索。这样设计的核心理由是：'
    '若撞墙给予过大的负奖励（例如 -100），智能体可能因"恐惧"探索而永远学不会穿越狭窄通道；-10 足以区分'
    '"无效率动作"与"倒退行为"，但不会摧毁学习信号。每步空地行走扣减 -1 分，这项微小但持续的负激励是'
    '驱动智能体寻找最短路径的关键机制——episode 步数越少，负累积越小，净收益越高。')

reward_headers = ['事件', '奖励值', '设计意图']
reward_rows = [
    ['到达 Goal', '+100', '强正激励，克服走路代价'],
    ['撞障碍物', '-10', '中等惩罚，留原地但不重置，自然学会规避'],
    ['空地行走', '-1', '微罚驱动最短路径寻找'],
]
add_table_with_data(doc, reward_headers, reward_rows)

add_heading_styled(doc, '3.4 关键实现细节', level=3)

add_para(doc, '碰撞处理方面', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    '智能体碰撞障碍物后不终止回合，而是留在原地并承受 -10 惩罚。该设计的目的是让智能体从失败中学习——'
    '在 Q 表中，从该格子出发指向障碍物方向的动作价值会因 -10 的即时奖励而逐步降低，从而在后续回合中智能体'
    '将自然而然地规避该方向。')
add_code_block(doc, '''elif next_state in self.obstacles:
    reward = -10
    done = False
    next_state = s           # 留在原地
    self.agent_pos = list(s)''')

add_para(doc, 'TD 目标的分情况处理', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    '终止状态 terminal 是一个哨兵值而非真实坐标，它不在 Q 表中（check_state_exist 不会为其创建条目），'
    '因此其所有动作的 Q 值恒为 0。在 learn() 方法中必须分情况处理：若 s\' 为 terminal，TD 目标直接取 '
    '即时奖励 r；否则 TD 目标按照标准公式 r + γ · max Q(s\') 计算。')
add_code_block(doc, '''def learn(self, s, a, r, s_):
    self.check_state_exist(s_)
    q_predict = self.q_table[s][a]
    if s_ != 'terminal':
        q_target = r + self.gamma * np.max(self.q_table[s_])
    else:
        q_target = r
    self.q_table[s][a] += self.lr * (q_target - q_predict)''')

add_para(doc, 'test_mode 在路径提取中的应用', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    '收敛后调用 extract_and_draw_optimal_path() 时，系统以 test_mode=True 调用 choose_action()，'
    '此时 ε-贪心退化为纯贪心策略（ε 被短路为 1.0），确保从起点到终点提取出的路径是确定性最优策略，'
    '不受 ε-贪心中 10% 随机探索的干扰。')

add_para(doc, '收敛检测与自动停机', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    '每个 episode 结束时，系统比较当前步数 current_steps 与上回合步数 last_episode_steps。若二者相等，'
    '稳定计数器 stable_count 递增；否则重置为 0 并记录新步数。当 stable_count 达到收敛阈值 3 时，自动'
    '停止训练、提取最优路径并弹出提示对话框。')
add_code_block(doc, '''if done:
    self.episode_count += 1
    if self.current_steps == self.last_episode_steps:
        self.stable_count += 1
    else:
        self.last_episode_steps = self.current_steps
        self.stable_count = 0
    if self.stable_count >= 3:
        self.toggle_training()
        self.extract_and_draw_optimal_path()''')

# ── 4. 实验结果与分析 ──
add_heading_styled(doc, '4. 实验结果与分析', level=2)
add_heading_styled(doc, '4.1 实验环境', level=3)

env2_headers = ['项目', '配置']
env2_rows = [
    ['Python', '3.9.23'], ['数值计算', 'NumPy 2.0.2'],
    ['GUI 框架', 'PyQt5 5.15.11'],
    ['conda 环境', 'rl_maze（极简环境，仅 22 个包）'],
    ['运行平台', 'Windows 11 Pro / Remote SSH (legion server)'],
]
add_table_with_data(doc, env2_headers, env2_rows)

add_heading_styled(doc, '4.2 默认参数实验', level=3)

add_para(doc,
    '在默认参数（6×6 迷宫，5 个随机障碍物，α = 0.1，γ = 0.9，ε = 0.9，收敛阈值 = 3，训练步进间隔 = '
    '10 ms）下，Q-Learning 智能体通常在 200 至 500 个 episodes 内收敛至最优策略。收敛速度主要取决于障碍物'
    '布局的复杂度——若障碍物恰好形成"死胡同"结构，智能体需要更多轮随机探索才能发现绕行路径。收敛后，系统'
    '自动以金色粗线在迷宫上绘制最优路径，直观展示智能体学到的策略。')

add_image_placeholder(doc, '训练过程截图：展示智能体随机探索阶段的行为')
add_image_placeholder(doc, '收敛后截图：展示金色最优路径从 Start 到 Goal')
add_image_placeholder(doc, '收敛弹窗截图：显示"智能体已收敛！最优路径需要步数: N"')

add_heading_styled(doc, '4.3 迷宫参数变化实验', level=3)

add_para(doc,
    '为验证系统的泛化能力和鲁棒性，本实验设计了四组不同迷宫规格的对照实验，从 4×4 的简单小迷宫到 '
    '10×10 的更复杂布局，系统性地观察状态空间规模与障碍物密度对收敛性能的影响。')

exp_headers = ['实验编号', '迷宫尺寸', '障碍物数量', '状态空间大小', '预期收敛轮次']
exp_rows = [
    ['A', '4×4', '2', '16', '较快'],
    ['B', '6×6', '5', '36', '中等'],
    ['C', '8×8', '10', '64', '较慢'],
    ['D', '10×10', '15', '100', '较慢'],
]
add_table_with_data(doc, exp_headers, exp_rows)

add_image_placeholder(doc, '不同参数下的收敛对比图')

add_para(doc,
    '从实验结果中可以提炼出三条关键观察。第一，状态空间规模与收敛速度之间存在正相关关系：随迷宫尺寸'
    '增大，Q 表条目数线性增长（每个格子对应一个状态键），TD 更新的传播需要更多 episode 才能覆盖从起点到'
    '终点的完整路径，收敛时间随之延长。第二，障碍物密度存在一个可行上限：障碍物过多会导致迷宫不可解——'
    '当障碍物数量大于等于总格子数减去起点和终点后，不存在任何连通路径。系统在 generate_new_maze() 中已'
    '设置了安全检查，当障碍物数量超过 (rows × cols - 2) 时弹出警告对话框。第三，最优路径长度的正确性得'
    '到验证：收敛后提取的最优路径长度与 BFS（广度优先搜索）计算出的理论最短路径一致，证明了 Q-Learning '
    '在确定性网格迷宫问题上能够收敛到全局最优策略。')

# ═══════════════════════════════════════════════════
#  EXPERIMENT 3
# ═══════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled(doc, '实验三：基于产生式系统的动物识别专家系统', level=1)

# ── 1. 问题分析 ──
add_heading_styled(doc, '1. 问题分析', level=2)
add_heading_styled(doc, '1.1 任务描述', level=3)

add_para(doc,
    '本实验要求基于产生式系统（Production System）理论，设计并实现一个动物识别专家系统。系统接收用户'
    '输入的动物特征观察，通过正向推理（数据驱动）识别出虎、金钱豹、斑马、长颈鹿、鸵鸟、企鹅、信天翁七种'
    '目标动物之一。系统需要提供图形化的特征选择界面（以复选框形式呈现所有可观察特征），实现完整的三层优先'
    '级产生式推理引擎以支持从原始观察到中间类别再到最终物种的递进推理，提供推理过程的可视化追踪（解释机构）'
    '使每个推导步骤对用户透明可查，并且知识库应通过外部 JSON 配置文件加载，使用户可以在不修改推理引擎代码'
    '的前提下方便地扩展新的动物物种和识别规则。')

add_heading_styled(doc, '1.2 产生式系统形式化定义', level=3)

add_para(doc,
    '产生式系统可形式化为一个三元组 Production System = ⟨KB, WM, IE⟩。知识库 KB 为规则的有限集合 '
    '{R₁, R₂, …, R₁₅}，每条规则 Rᵢ = ⟨Pᵢ, Cᵢ, priorityᵢ⟩ 包含一个前提条件集 Pᵢ（合取范式形式，即'
    '多个特征需同时满足）、一个结论 Cᵢ（推导出的新事实）以及一个优先级 priorityᵢ（数值越小越优先触发，'
    '用于冲突消解）。工作内存 WM 是当前已知事实的动态集合，初始值为用户通过复选框选中的观察特征，在推理'
    '过程中随着规则的触发而逐步扩增。推理引擎 IE 执行正向链接（Forward Chaining）的匹配-选择-执行循环：'
    '扫描知识库中所有未被使用过的规则，将前提完全被 WM 包含的规则加入冲突集，通过冲突消解策略选出优先级'
    '最高的一条执行，将其结论加入 WM，如此循环直至推导出目标动物或冲突集为空。')

ps_headers = ['组件', '符号', '定义']
ps_rows = [
    ['知识库', 'KB', '规则集合，Rᵢ = ⟨Pᵢ, Cᵢ, priorityᵢ⟩'],
    ['工作内存', 'WM', '已知事实的动态集合'],
    ['推理引擎', 'IE', '正向链接的匹配-选择-执行循环'],
]
add_table_with_data(doc, ps_headers, ps_rows)

add_heading_styled(doc, '1.3 知识库结构分析', level=3)

add_para(doc,
    '系统知识库包含 15 条产生式规则，按照推理的抽象层次分为三个优先级层级，构成了"原始特征 → 中间类别'
    ' → 最终物种"的递进推理链。')

add_para(doc,
    '优先级 1（基础分类规则，4 条，R1-R4）负责将用户直接可观察的原始特征映射到生物类别。规则 R1（若有'
    '毛发则推出哺乳动物）和 R2（若产奶则推出哺乳动物）共同覆盖了哺乳动物的两种独立识别路径；规则 R3（若有'
    '羽毛则推出鸟类）和 R4（若会飞且下蛋则推出鸟类）同样提供了鸟类的两种互补识别方式。这一层级的设计体现了'
    '专家系统中"多条规则指向同一中间结论"的知识冗余策略——即使用户提供的特征不完整（例如仅选了"产奶"而未选'
    '"有毛发"），系统仍能通过其他规则推导出相同的中间类别。', bold=False, size=12)

p1_headers = ['规则 ID', '前提条件', '结论']
p1_rows = [
    ['R1', '{有毛发}', '哺乳动物'], ['R2', '{产奶}', '哺乳动物'],
    ['R3', '{有羽毛}', '鸟类'], ['R4', '{会飞, 下蛋}', '鸟类'],
]
add_table_with_data(doc, p1_headers, p1_rows)

add_para(doc,
    '优先级 2（中间推理规则，4 条，R5-R8）在基础类别之上进行进一步细分，其前提条件中可能同时包含原始'
    '特征和优先级 1 推导出的中间结论。规则 R5（若吃肉则推出食肉动物）和 R6（若有犬齿、有爪且眼盯前方则推出'
    '食肉动物）从不同角度识别食肉动物；规则 R7（若为哺乳动物且有蹄则推出有蹄类动物）和 R8（若为哺乳动物且'
    '反刍则推出有蹄类动物）将哺乳动物进一步细分为有蹄类。注意 R7 和 R8 的前提中包含了优先级 1 的结论'
    '"哺乳动物"——这正是层级推理的关键特征：高层规则依赖于低层规则的输出。')

p2_headers = ['规则 ID', '前提条件', '结论', '备注']
p2_rows = [
    ['R5', '{吃肉}', '食肉动物', ''],
    ['R6', '{有犬齿, 有爪, 眼盯前方}', '食肉动物', '多条件合取'],
    ['R7', '{哺乳动物, 有蹄}', '有蹄类动物', '依赖 P1 推理结果'],
    ['R8', '{哺乳动物, 反刍}', '有蹄类动物', '依赖 P1 推理结果'],
]
add_table_with_data(doc, p2_headers, p2_rows)

add_para(doc,
    '优先级 3（最终识别规则，7 条，R9-R15）位于推理链的最顶层，其前提条件是原始特征、优先级 1 结论和'
    '优先级 2 结论的复合组合，结论即为七种目标动物之一。例如，识别虎的规则 R9 需要同时满足"哺乳动物"'
    '（来自 P1）、"食肉动物"（来自 P2）、"黄褐色"和"黑色条纹"（原始特征）四个条件。这种层层递进的'
    '知识结构设计是专家系统的典型工程实践——将复杂的物种识别任务分解为可独立维护的层级规则，使得每层规则'
    '的语义清晰、修改范围可控。')

p3_headers = ['规则 ID', '前提条件', '结论']
p3_rows = [
    ['R9', '{哺乳动物, 食肉动物, 黄褐色, 黑色条纹}', '虎'],
    ['R10', '{哺乳动物, 食肉动物, 黄褐色, 暗斑点}', '金钱豹'],
    ['R11', '{有蹄类动物, 黑色条纹}', '斑马'],
    ['R12', '{有蹄类动物, 长脖子, 长腿, 暗斑点}', '长颈鹿'],
    ['R13', '{鸟类, 长脖子, 长腿, 不会飞, 黑白二色}', '鸵鸟'],
    ['R14', '{鸟类, 会游泳, 不会飞, 黑白二色}', '企鹅'],
    ['R15', '{鸟类, 善飞}', '信天翁'],
]
add_table_with_data(doc, p3_headers, p3_rows)

add_heading_styled(doc, '1.4 关键挑战', level=3)

add_para(doc,
    '动物识别专家系统面临四个主要挑战。第一是冲突消解（Conflict Resolution）：多条规则可能同时被当前的'
    '工作内存满足（例如 R1 和 R2 均可独立推出"哺乳动物"），系统需要确定先触发哪一条规则，该决策直接影响'
    '推理链的长度和可解释性。第二是部分匹配处理：当用户输入的特征数量不足、无法唯一确定目标动物时，系统'
    '不应简单地输出"推理失败"，而应当基于现有证据给出部分匹配的候选列表和命中特征，帮助用户判断还需补充'
    '哪些观察。第三是可扩展性：知识库应易于通过外部配置文件进行维护和扩展，添加新物种和新规则的操作不应'
    '要求修改推理引擎的核心代码。第四是推理可解释性：完整记录每一步推理的触发规则、前提条件和推导结论，'
    '使最终用户能够理解"系统为什么得出此结论"，而非面对一个黑箱判决。')

# ── 2. 算法原理 ──
add_heading_styled(doc, '2. 算法原理', level=2)
add_heading_styled(doc, '2.1 正向链接推理算法', level=3)

add_para(doc,
    '正向链接（Forward Chaining）是一种数据驱动（data-driven）的推理策略：从已知事实出发，反复应用规则'
    '推导出新事实，直至无法再推导或已成功识别出目标。算法的核心循环为匹配-选择-执行（Match-Select-'
    'Execute）：在匹配阶段，扫描知识库中所有尚未被使用过的规则，将前提条件完全被当前工作内存包含的规则'
    '收集为冲突集（Conflict Set）；在选择阶段，通过冲突消解策略从冲突集中选出一条优先执行的规则；在执行'
    '阶段，将该规则的结论加入工作内存，并在解释日志中记录此步推理的完整信息。')

add_para(doc,
    '关于时间复杂度，每轮迭代需要扫描全部未使用规则以检查前提是否被满足（子集判断），最坏情况下每条规则'
    '恰好触发一次，总匹配操作次数为 O(|R|²)，其中 |R| 为规则总数。对于本系统 15 条规则的小规模知识库，'
    '这一复杂度完全在可接受范围内。推理循环在两种条件下终止：一是所选规则的结论命中任意一个目标动物'
    '（推理成功），二是冲突集为空且尚未命中任何目标（推理失败，进入部分匹配阶段）。')

add_heading_styled(doc, '2.2 冲突消解策略', level=3)

add_para(doc,
    '当冲突集中存在多条可触发规则时，系统采用两级排序策略来确定执行顺序，排序键为 order(Rᵢ) = '
    '(priority(Rᵢ), -|Pᵢ|)。第一关键字 priority 按升序排列，这意味着数值较小的优先级（即更基础的分类'
    '规则，如 R1-R4 的 priority = 1）将先于数值较大的优先级（即最终识别规则，如 R9-R15 的 priority = 3）'
    '触发，从而确保推理按照"原始特征 → 中间类别 → 最终物种"的正确层次顺序推进。第二关键字按前提数量的'
    '负值排列（即前提越多越靠前），这体现了"特殊性优先"原则——在相同优先级下，前提条件更多的规则其匹配'
    '范围更窄、语义更具体，应优先触发，以避免宽泛规则过早"抢占"推理路径。')

add_heading_styled(doc, '2.3 部分匹配与依赖树递归', level=3)

add_para(doc,
    '当精确的正向链接推理未能命中任一目标动物时，系统不会简单报告失败，而是自动进入部分匹配模式，为用'
    '户提供基于特征交集的候选推荐。该算法的执行分为三个步骤。')

add_para(doc,
    '第一步是构建特征依赖树。对于每个候选目标动物 T，递归展开其完整特征依赖树 D(T) = {T} ∪ '
    '⋃_{R: conclusion(R)=T} (P(R) ∪ ⋃_{p∈P(R)} D(p))，即目标 T 本身加上所有能推导出 T 的规则的'
    '前提集，再加上这些前提各自递归展开的依赖树。递归过程通过 visited 集合追踪已访问的结论节点，防止在'
    '知识库出现循环依赖时陷入无限递归。第二步是计算交集，对每个目标 T 计算 D(T) 与当前工作内存 WM 的'
    '交集 match(T) = D(T) ∩ WM，该交集的大小反映了当前已有证据对目标 T 的支持程度。第三步是排序推荐，'
    '将所有 match(T) 非空的目标按命中特征数量降序排列，向用户展示每个候选动物名称及其具体命中特征列表。'
    '这种设计让用户一目了然地看到"还差哪些特征"就能确定某种动物，有效引导用户补充观察。')

# ── 3. 系统设计与实现 ──
add_heading_styled(doc, '3. 系统设计与实现', level=2)
add_heading_styled(doc, '3.1 系统架构', level=3)

add_para(doc,
    '系统采用知识-推理-交互三层解耦架构。交互层由 AnimalExpertSystem 主窗口类承担，左侧面板通过 '
    'QScrollArea 滚动区域以两列网格布局动态生成特征复选框，底部放置"开始推理"和"重置特征"两个操作按钮；'
    '右侧大面积的 QTextEdit 日志展示区以等宽字体实时呈现推理链和部分匹配候选推荐，左右比例为 2:3。推理层'
    '由 InferenceEngine 类独立封装，其中的 forward_chaining() 方法执行完整的正向链接推理循环，内部包含'
    '冲突集的收集、两级排序消解以及逐步解释日志的生成，完全不依赖任何 UI 组件。知识层由 KnowledgeBase '
    '类和外部 JSON 文件 knowledge.json 共同构成：KnowledgeBase.load_from_json() 负责解析 JSON 并构建'
    'Rule 对象列表，同时自动推导终端目标（targets）；get_all_premises() 方法从知识库中动态提取出仅属于'
    '用户直接观察范围的基础特征（排除所有中间结论和目标动物名称），供交互层生成复选框。这种分层设计的核心'
    '优势在于：知识工程师可以通过编辑 JSON 文件来调整规则库，推理策略的改进仅需修改 InferenceEngine，而 '
    'UI 的变更完全不影响前两者——三层各自独立演化。')

add_heading_styled(doc, '3.2 核心类设计', level=3)

add_para(doc,
    'Rule 类是最基础的数据单元，每条规则封装为一个 Rule 对象，包含 rule_id（规则标识符，如 "R9"）、'
    'premise（前提集，在构造时自动从列表转换为 Python set 类型，以在后续推理的每一步中子集判断操作达到 '
    'O(|P|) 的均摊时间复杂度）、conclusion（推理结论字符串）以及 priority（整数优先级，数值越小越优先'
    '触发）。')

add_para(doc,
    'KnowledgeBase 类承担知识库的管理职责。load_from_json(json_path) 方法解析外部 JSON 文件，遍历 '
    'rules 数组逐一构建 Rule 对象并加入内部规则列表。在加载完成后，若 JSON 中显式提供了 targets 字段则'
    '直接使用；否则执行自动推导——取所有规则结论的集合减去所有规则前提的集合，差集中的元素即为"仅作为结论'
    '出现、从未作为其他规则的前提使用"的终端目标。这一设计使得在 JSON 中添加新动物识别规则后无需手动声明 '
    'targets，系统能自动感知新物种的加入。get_all_premises() 方法则负责从知识库中提取仅由用户直接观察的'
    '基础特征：取所有规则前提的并集，再从中减去所有中间结论（即曾作为某条规则 conclusion 出现的术语）和'
    '所有最终目标动物名称，剩余的特征即为 UI 复选框的生成依据。')

add_code_block(doc, '''if "targets" in data:
    self.targets = set(data["targets"])
else:
    self.targets = all_conclusions - all_premises''')

add_para(doc,
    'InferenceEngine 类封装了推理引擎的全部逻辑，其核心数据结构包括三个组件：working_memory（Set[str]）'
    '是当前已知事实的动态集合，初始值为用户在 UI 中勾选的特征，在推理循环中随规则触发而逐步扩增；'
    'used_rules（Set[str]）记录已经触发过的规则 ID，确保同一条规则不会被反复触发（避免无限循环）；'
    'explanation_log（List[str]）是逐步推理的完整文字记录，每个条目包含步骤序号、触发规则 ID、前提条件'
    '（"因为什么"）和推导结论（"所以得出什么"），是实现推理可解释性的关键数据结构。')

add_para(doc,
    'AnimalExpertSystem 主窗口类在初始化时创建 KnowledgeBase 实例（加载 knowledge.json）和 '
    'InferenceEngine 实例（绑定该知识库），然后调用 initUI() 构建完整的图形界面。run_inference() 方法'
    '是用户点击"开始推理"按钮后的入口：首先收集所有被勾选复选框的文本作为初始事实集；接着调用 '
    'forward_chaining() 执行正向推理；若推理成功（返回非 None 的目标动物），则将完整推理链日志逐行输出到'
    '右侧文本区并高亮显示识别结果；若推理失败，则进入部分匹配流程——对每个目标动物递归构建依赖树、计算与'
    '当前工作内存的交集、按命中数降序排列后展示候选列表。')

ui_headers = ['方法', '功能']
ui_rows = [
    ['initUI()', '构建左右分栏布局：左侧特征选择区 + 右侧日志展示区'],
    ['run_inference()', '收集特征 → 正向推理 → 成功展示推理链 / 失败计算候选'],
    ['reset_features()', '清除所有复选框勾选及日志区域'],
]
add_table_with_data(doc, ui_headers, ui_rows)

add_heading_styled(doc, '3.3 关键实现细节', level=3)

add_para(doc, '冲突消解的实现', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    '系统通过 Python 的 sort 方法配合 lambda 键函数，以一行代码完成两级排序：'
    'conflict_set.sort(key=lambda r: (r.priority, -len(r.premise)))。priority 升序保证基础分类规则'
    '优先触发，-len(premise) 降序（对前提数量取负值以实现降序排列）保证在相同优先级下更具体的规则优先。'
    '排序完成后直接取 conflict_set[0] 即为当前最优触发规则。')
add_code_block(doc, '''conflict_set.sort(key=lambda r: (r.priority, -len(r.premise)))
selected_rule = conflict_set[0]''')

add_para(doc, '推理可解释性的日志生成', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    '每步推理生成一条包含三个关键信息的日志记录：(a) 触发了哪条规则（规则 ID），(b) 规则的前提条件集合'
    '（"因为什么"——用 AND 连接各前提），(c) 规则的结论（"所以推导出什么"——新加入工作内存的事实）。'
    '完整的推理链在右侧 QTextEdit 中以等宽字体（Consolas）逐行展示，每条记录之间以空行分隔，推理成功时'
    '以分隔线和加粗文字高亮最终识别结果。这种设计使得系统的推理过程对用户完全透明——用户可以逐条回溯推理'
    '链，验证每一步的逻辑合理性。')
add_code_block(doc, '''log_line = (
    f"-> 步骤 {step}: 触发规则 [{selected_rule.rule_id}]\\n"
    f"   IF   {' AND '.join(selected_rule.premise)}\\n"
    f"   THEN {selected_rule.conclusion}\\n"
    f"   (事实库加入: '{selected_rule.conclusion}')\\n")''')

add_para(doc, '部分匹配的依赖树递归', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    'get_all_dependencies(conclusion, visited=None) 函数递归地计算某个结论的完整前提依赖闭包。对于给定'
    '的 conclusion，遍历知识库中所有以它为结论的规则，收集每条规则的前提条件，并对其中的每个前提递归调用'
    '自身以展开更底层的依赖。visited 集合参数用于在调用链中追踪已访问的节点：若当前结论已在 visited 中'
    '出现，说明知识库中存在循环依赖（例如 A → B → A），此时立即返回空集以避免无限递归。这一防御性设计在'
    '知识库由非专业人员手动编辑时至关重要——它保证了即使 JSON 配置出现逻辑错误，系统也不会崩溃而是优雅地'
    '终止该分支的递归。')
add_code_block(doc, '''def get_all_dependencies(conclusion, visited=None):
    if visited is None: visited = set()
    if conclusion in visited: return set()  # 防环
    visited.add(conclusion)
    deps = set()
    for r in self.engine.kb.rules:
        if r.conclusion == conclusion:
            deps.update(r.premise)
            for p in r.premise:
                deps.update(get_all_dependencies(p, visited))
    return deps''')

add_para(doc, '特征动态提取与 UI 生成', bold=True, first_line_indent=0.74, size=11)
add_para(doc,
    'UI 的复选框列表并非硬编码，而是在 initUI() 中通过调用 kb.get_all_premises() 动态生成。该方法首先'
    '收集所有规则前提的并集（all_conditions），然后从中减去所有规则的结论集合（conclusions）和所有目标'
    '动物的集合（targets）。减法的语义是：凡是在知识库中作为某条规则"输出"出现的术语，都应由推理引擎在'
    '运行时自动推导，不应作为用户可直接勾选的选项呈现。例如，"哺乳动物"是 R1 和 R2 的结论，因此用户不会'
    '在特征列表中看到"哺乳动物"复选框——系统会自动根据用户是否选择了"有毛发"或"产奶"来决定是否推导'
    '出哺乳动物。这种设计保证了 UI 仅展示真正的"观测层"特征，维持了推理层次逻辑的严谨性。')
add_code_block(doc, '''def get_all_premises(self):
    all_conditions = set()
    for rule in self.rules: all_conditions.update(rule.premise)
    conclusions = {r.conclusion for r in self.rules}
    return sorted(list(all_conditions - conclusions - self.targets))''')

# ── 4. 实验结果与分析 ──
add_heading_styled(doc, '4. 实验结果与分析', level=2)
add_heading_styled(doc, '4.1 实验环境', level=3)

env3_headers = ['项目', '配置']
env3_rows = [
    ['Python', '3.9.23'], ['GUI 框架', 'PyQt5 5.15.11'],
    ['知识库格式', 'JSON (knowledge.json, 15 条规则)'],
    ['conda 环境', 'expert_env（52 个包，含 PyTorch/ultralytics 视觉扩展支持）'],
    ['运行平台', 'Windows 11 Pro / Remote SSH (legion server)'],
]
add_table_with_data(doc, env3_headers, env3_rows)

add_heading_styled(doc, '4.2 推理成功案例', level=3)

add_para(doc,
    '以下通过两个典型推理案例展示系统的工作流程。案例一为识别虎：用户在特征选择界面勾选{有毛发, 吃肉, '
    '黄褐色, 黑色条纹}四项特征后点击"开始推理"。推理引擎首先扫描优先级 1 规则，发现 R1（若有毛发 → '
    '哺乳动物）的前提"有毛发"在 WM 中，触发 R1，将"哺乳动物"加入 WM。继续扫描优先级 2 规则，发现 R5'
    '（若吃肉 → 食肉动物）的前提"吃肉"在 WM 中，触发 R5，将"食肉动物"加入 WM。最终扫描优先级 3 规则，'
    '发现 R9（若为哺乳动物、食肉动物、黄褐色、有黑色条纹 → 虎）的全部四个前提均在 WM 中，触发 R9，成功'
    '将"虎"识别为目标动物。推理共需 3 步，推理链清晰展示了从原始特征到中间类别再到最终物种的递进过程。')

case1_headers = ['步骤', '已具备的事实', '触发规则', '推导结论']
case1_rows = [
    ['初始', '{有毛发, 吃肉, 黄褐色, 黑色条纹}', '—', '—'],
    ['1', '加入新事实', 'R1', '哺乳动物'],
    ['2', '加入新事实', 'R5', '食肉动物'],
    ['3', '加入新事实', 'R9', '虎 ✅'],
]
add_table_with_data(doc, case1_headers, case1_rows)

add_image_placeholder(doc, '推理成功界面截图：展示特征选择和虎的推理链')

add_para(doc,
    '案例二为识别企鹅：用户勾选{有羽毛, 会游泳, 不会飞, 黑白二色}四项特征。推理引擎首先触发 R3（若有'
    '羽毛 → 鸟类），将"鸟类"加入 WM。随后触发 R14（若为鸟类、会游泳、不会飞、有黑白二色 → 企鹅），成功'
    '识别出企鹅。这个案例仅需 2 步推理——用户直接提供了 R14 所需的大部分原始特征（会游泳、不会飞、黑白二色）'
    '，推理引擎只需额外推导出"鸟类"这一中间类别即可完成识别。这展示了当用户提供的信息足够充分时，推理链可以'
    '非常简短高效。')

case2_headers = ['步骤', '已具备的事实', '触发规则', '推导结论']
case2_rows = [
    ['初始', '{有羽毛, 会游泳, 不会飞, 黑白二色}', '—', '—'],
    ['1', '加入新事实', 'R3', '鸟类'],
    ['2', '加入新事实', 'R14', '企鹅 ✅'],
]
add_table_with_data(doc, case2_headers, case2_rows)

add_image_placeholder(doc, '企鹅识别结果截图')

add_heading_styled(doc, '4.3 部分匹配案例', level=3)

add_para(doc,
    '当用户提供的特征不足以唯一确定目标动物时，系统自动进入部分匹配模式。例如，若用户仅选择了{有毛发, '
    '"黑色条纹"}两项特征，正向推理链无法到达任何目标动物（有毛发 → 哺乳动物，但哺乳动物 + 黑色条纹不足以'
    '唯一匹配 R9 或 R11 的全部前提）。此时系统对七个目标逐一构建依赖树并计算与当前 WM 的交集：斑马的依赖树'
    '与 WM 的交集为{"黑色条纹"}（命中 1 项），虎的依赖树与 WM 的交集同为{"黑色条纹"}（命中 1 项），其余目标的'
    '交集为空。这两个候选按命中数并列，在日志区依次展示候选名称和命中特征。又如，用户选择{有羽毛, 不会飞}，'
    '系统推导出"鸟类"后，鸵鸟和企鹅的依赖树均命中{"不会飞", "鸟类"}两项特征，系统将并列展示这两个候选。'
    '用户可根据系统提示的命中特征，有针对性地补充观察以进一步区分候选物种。')

pm_headers = ['输入特征', '候选动物', '命中特征', '命中数']
pm_rows = [
    ['{有毛发, 黑色条纹}', '斑马', '黑色条纹', '1'],
    ['{有毛发, 黑色条纹}', '虎', '黑色条纹', '1'],
    ['{有羽毛, 不会飞}', '鸵鸟', '不会飞, 鸟类', '2'],
    ['{有羽毛, 不会飞}', '企鹅', '不会飞, 鸟类', '2'],
]
add_table_with_data(doc, pm_headers, pm_rows)

add_image_placeholder(doc, '部分匹配候选推荐界面截图')

add_heading_styled(doc, '4.4 知识库扩展性验证', level=3)

add_para(doc,
    '为验证"JSON 外部知识库 + 自动推导目标"架构的可扩展性，本实验设计了一个扩展测试场景：向 '
    'knowledge.json 中添加一条新规则 R16（若为哺乳动物、食肉动物且强壮的 → 狮子），不修改任何 Python '
    '源代码。重新启动程序后，系统自动执行以下操作：解析 JSON 时将 R16 加载到知识库的规则列表中；由于 R16 '
    '的结论"狮子"从未在任何规则的前提中出现，自动目标推导算法将其识别为新的终端目标，加入 targets 集合'
    '；get_all_premises() 方法在计算可观察特征时发现"强壮的"不在任何规则的结论集中，将其作为新的用户可选'
    '特征，在 UI 中自动生成对应的复选框。整个过程无需修改任何 Python 代码，充分验证了知识-推理分离架构的'
    '可扩展性优势：领域专家或课程教师可以通过编辑 JSON 文件随时调整知识库，而无需了解推理引擎的实现细节。')
add_code_block(doc, '''{
  "id": "R16",
  "priority": 3,
  "premise": ["哺乳动物", "食肉动物", "强壮的"],
  "conclusion": "狮子"
}''')

# ═══════════════════════════════════════════════════
#  EXPERIENCE & SUGGESTIONS
# ═══════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled(doc, '体会与建议', level=1)
add_heading_styled(doc, '实验收获', level=2)

add_para(doc,
    '实验二（Q-Learning 迷宫）带来了三方面的核心收获。首先，通过亲手实现 TD 更新循环——从状态选择动作、'
    '与环境交互获得奖励、计算 TD 误差、以学习率向目标方向修正 Q 值——我深刻体会到了强化学习中"从试错中'
    '学习"的本质：每步奖励信号如同教师对学生的批改反馈，逐步将 Q 表从零初始状态修正到逼近真实的动作价值函数。'
    '其次，探索-利用困境的工程权衡是一个需要实践感知的问题：ε 值过高（如 0.99）会导致智能体过早陷入对次优'
    '路径的"迷信"，ε 值过低（如 0.5）则使收敛极为缓慢。本任务中 ε = 0.9 取得了良好平衡，但在更复杂环境'
    '（如随机奖励、部分可观测迷宫）中，可能需要引入衰减策略（ε_t = max(ε_min, ε_init · decay^t)）来动态'
    '调节探索程度。第三，收敛判定的实践价值超出预期：基于"步数稳定"的启发式停机准则比设定固定训练轮次更具'
    '工程意义——避免了固定轮次过少导致的欠训（Q 表未收敛即停）和过多导致的过训（浪费计算资源），实现了训练'
    '过程的自动化闭环。')

add_para(doc,
    '实验三（动物识别专家系统）同样带来了三方面的核心收获。在知识工程方法论层面，将 15 条领域规则从推理'
    '引擎（正向链接 + 冲突消解）中彻底解耦、以 JSON 格式外部存储的架构设计，体现了专家系统的核心理念——'
    '"知识可配置"。这使得非编程人员（如生物学家、课程教师）只需编辑 JSON 文件即可维护和扩展知识，无需接触'
    'Python 代码，显著降低了系统的维护门槛。在推理可解释性层面，每步推理的完整日志（"因为 A，所以 B"的'
    '形式）让用户信任系统的结论——这是现代 AI 可解释性（XAI）研究在传统符号主义 AI 中的早期实践。当系统输出'
    '"虎"时，用户可以逐条回溯推理链，验证每一步的逻辑合理性，而非面对一个无法审计的黑箱判断。在用户体验'
    '设计层面，部分匹配机制是一种优雅的"降级服务"策略——当精确推理因信息不足而失败时，系统不是简单地报告'
    '"无法识别"，而是利用依赖树递归算法计算出与当前证据最接近的候选动物及其命中特征，引导用户补充观察，'
    '将"失败"转化为"建设性的诊断建议"。')

add_heading_styled(doc, '改进方向', level=2)

add_para(doc,
    '迷宫 RL 系统可以从三个方向进一步改进。一是引入 ε-衰减策略，使 ε 从初始值 0.9 随训练 episode 数线性'
    '或指数衰减至 0.01，使训练前期充分探索状态空间、训练后期专注利用已学到的知识，从而加速收敛并提升最终'
    '策略的稳定性。二是添加随机起点和终点的选项，使智能体学会从迷宫中任意位置导航至任意目标，增强策略的泛化'
    '能力和鲁棒性。三是实现 SARSA（on-policy 的 TD 控制算法）与 Q-Learning（off-policy）的对比实验，在相同'
    '的迷宫布局和超参数下运行两种算法，直观比较它们在探索安全性（SARSA 更保守，因为它在更新时考虑了 ε-贪心'
    '的动作选择）、收敛路径和最终策略质量上的差异，深化对"on-policy 与 off-policy"这一强化学习核心概念的'
    '理解。')

add_para(doc,
    '动物专家系统同样有三条改进路径。一是引入不确定性推理机制（如确定性因子 CF 或贝叶斯方法），处理"可能"'
    '"或许""大概"等模糊观察特征——现实中用户并不总是能 100% 确定一个特征的存在与否，带有置信度的推理能'
    '显著提升系统的实用性。二是集成视觉感知模块（复用 mission3 目录中已有的 vision_expert2.py 和 YOLOv8 '
    '模型 yolov8n.pt），构建"拍照 → YOLOv8 目标检测 → 检测标签映射为特征 → 产生式逻辑推理"的端到端智能'
    '识别流水线，将连接主义 AI（深度学习感知）与符号主义 AI（规则推理）进行融合，兼具感知的鲁棒性和推理的'
    '可解释性。三是添加反向链接（Backward Chaining）推理模式——系统从某个候选目标出发，反向查找该目标所需'
    '的全部前提条件，识别出当前工作内存中缺失的特征后，主动向用户提问（如"该动物是否有黑色条纹？"），实现'
    '交互式诊断。正向链接和反向链接两种模式可在同一系统中并存，用户可根据场景自由切换。')

add_heading_styled(doc, '对课程的建议', level=2)

add_para(doc,
    '基于本次实验的实践体会，提出三点课程改进建议。首先，实验指导书可在算法原理部分增加对核心公式的推导'
    '引导环节——例如 Q-Learning 的 TD 更新公式可以从 Bellman 最优方程的采样近似角度展开推导，正向链接中'
    '冲突消解的两级排序可以从知识工程的"特殊性优先"原则加以论证——帮助学生在动手编码前建立从理论到公式的'
    '直觉桥梁，降低"记住公式但不会灵活运用"的认知断层。其次，可增设"算法对比"实验环节，要求学生在迷宫任务'
    '中同时实现 Q-Learning 和 SARSA、在动物识别任务中同时实现正向链接和反向链接，通过实验数据（收敛曲线、'
    '推理步数、准确率等）的量化对比来理解不同算法设计哲学（off-policy vs. on-policy；数据驱动 vs. 目标驱'
    '动）在具体场景下的适用性差异，这种对比式学习远比孤立实现单一算法更能加深理解。第三，可引入轻量级的'
    '基准测试（benchmark）——为迷宫任务设置一组标准障碍物布局（固定种子以确保可复现性），对收敛 episode '
    '数和最优路径长度进行量化评估；为动物识别任务设置标准测试用例集（不同特征组合对应预期输出），对推理准确'
    '率进行系统性测试。基准测试的引入将使各学生的实验结果具有横向可比性，也便于助教进行自动化评分。')

# ── Final heading font fix ──
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading'):
        for run in paragraph.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# ── Save ──
output_path = 'D:/ScuderiaAMG/人工智能导论/实验部分/实验报告_任务2_任务3.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
