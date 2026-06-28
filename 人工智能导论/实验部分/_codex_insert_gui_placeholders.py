from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX = max(
    [p for p in ROOT.iterdir() if p.suffix.lower() == ".docx" and not p.name.startswith("~$")],
    key=lambda p: p.stat().st_mtime,
)


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=None, italic=None):
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_border(cell, val="single", sz="8", color="7F7F7F"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), sz)
        node.set(qn("w:color"), color)


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_row_height(row, cm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(int(cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")


def format_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    for run in p.runs:
        set_run_font(run, size=9)


def insert_gui_placeholder(anchor, caption, note):
    intro = anchor.insert_paragraph_before()
    intro.paragraph_format.first_line_indent = Cm(0.74)
    intro.paragraph_format.line_spacing = 1.25
    intro.paragraph_format.space_after = Pt(3)
    r = intro.add_run(note)
    set_run_font(r)

    table = anchor.insert_paragraph_before()._p
    # python-docx has no insert_table_before helper, so create at document end and move it.
    doc = anchor._parent
    placeholder = doc.add_table(rows=1, cols=1)
    placeholder.alignment = WD_TABLE_ALIGNMENT.CENTER
    placeholder.allow_autofit = False
    placeholder.rows[0].cells[0].width = Cm(13.5)
    set_row_height(placeholder.rows[0], 6.8)
    cell = placeholder.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)
    set_cell_shading(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run("此处插入最终图形化界面截图")
    set_run_font(run, east="宋体", west="Times New Roman", size=11, italic=True)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    anchor._p.addprevious(placeholder._tbl)

    cap = anchor.insert_paragraph_before(caption)
    format_caption(cap)


doc = Document(str(DOCX))
paragraphs = doc.paragraphs
mission2_anchor = None
mission3_anchor = None
for p in paragraphs:
    text = p.text.strip()
    if text.startswith("2.1.4"):
        mission2_anchor = p
    if text.startswith("2.2.4"):
        mission3_anchor = p

if mission2_anchor is None or mission3_anchor is None:
    raise RuntimeError("Could not locate Mission02 or Mission03 placeholder anchors.")

insert_gui_placeholder(
    mission2_anchor,
    "图 2-1  Mission02 Q-Learning 迷宫寻路系统最终图形化界面（待插入）",
    "为便于呈现最终运行效果，Mission02 在算法与界面调度讨论之后预留图形化界面截图位置。该图建议展示迷宫网格、起点、终点、障碍物、训练控制面板及收敛后的最优路径。",
)
insert_gui_placeholder(
    mission3_anchor,
    "图 2-2  Mission03 动物识别专家系统最终图形化界面（待插入）",
    "为便于呈现最终运行效果，Mission03 在知识库、推理链与视觉扩展讨论之后预留图形化界面截图位置。该图建议展示特征选择区域、推理按钮、解释日志和识别结果输出区域。",
)

doc.save(str(DOCX))
print(DOCX)
