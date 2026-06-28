from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX = max(
    [p for p in ROOT.iterdir() if p.suffix.lower() == ".docx" and not p.name.startswith("~$")],
    key=lambda p: p.stat().st_mtime,
)


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=None, italic=None):
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), west)
    rfonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_body(p):
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_run_font(run)


def format_heading3(p):
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        set_run_font(run, east="黑体", west="Arial", size=10.5, bold=True)


def format_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    for run in p.runs:
        set_run_font(run, size=9)


def set_cell_border(cell, val="single", sz="8", color="7F7F7F"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), sz)
        node.set(qn("w:color"), color)


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_row_height(row, cm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(int(cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_items_before(anchor, items):
    for text, style in items:
        p = anchor.insert_paragraph_before(text, style=style)
        if style == "Heading 3":
            format_heading3(p)
        else:
            format_body(p)


def insert_placeholder_before(doc, anchor, caption, note):
    p = anchor.insert_paragraph_before(note)
    format_body(p)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    row = table.rows[0]
    set_row_height(row, 6.8)
    cell = row.cells[0]
    cell.width = Cm(13.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)
    set_cell_shading(cell, "FFFFFF")
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(0)
    r = cp.add_run("此处插入最终图形化界面截图")
    set_run_font(r, size=11, italic=True)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    anchor._p.addprevious(table._tbl)

    cap = anchor.insert_paragraph_before(caption)
    format_caption(cap)


doc = Document(str(DOCX))

# Remove appendix heading, trailing blank paragraph, appendix table, and TOC entry.
for p in list(doc.paragraphs):
    if p.text.strip() == "附录A  源文件覆盖清单" or p.text.strip() == "附录A  源文件覆盖清单":
        remove_element(p._element)
    elif p.text.strip().startswith("附录A"):
        remove_element(p._element)

for table in list(doc.tables):
    try:
        first = table.cell(0, 0).text.strip()
    except Exception:
        first = ""
    if first == "源文件":
        remove_element(table._element)

for p in list(doc.paragraphs):
    if p.text.strip() == "附录A  源文件覆盖清单":
        remove_element(p._element)
    if p.text.strip() == "" and p._element.getprevious() is not None:
        prev_text = getattr(p._element.getprevious(), "text", None)

anchors = {}
for p in doc.paragraphs:
    t = p.text.strip()
    for key in [
        "1.2  Session02",
        "1.3  Session03",
        "1.4  Session04",
        "1.5  Session05",
        "2  最终四选二",
        "2.1.4",
        "2.2.4",
        "3  综合分析",
    ]:
        if t.startswith(key):
            anchors[key] = p

session1 = [
    ("1.1.1  数据特征与分类边界分析", "Heading 3"),
    ("Iris 数据集具有样本规模小、特征维度低和类别边界相对清晰的特点，适合作为支持向量机实验的基准数据。脚本将四个形态学特征读入后，主要使用前两个特征绘制二维分类平面。该处理牺牲了完整四维特征空间的信息，但换取了可视化解释能力，使分类边界、支持向量机核映射效果和样本分布之间的关系能够被直观观察。", "Normal"),
    ("采用 RBF 核函数的意义在于引入非线性相似度度量。若使用线性核，模型只能在原始特征平面上构造线性分割；RBF 核则依据样本间距离形成局部响应，能够拟合 Versicolor 与 Virginica 之间更复杂的过渡区域。C 与 gamma 的取值共同控制间隔约束和局部曲率：C 较大时模型更重视训练误差，gamma 较大时决策边界更容易贴近样本分布。当前参数更适合演示非线性边界，但也需要关注过拟合风险。", "Normal"),
    ("1.1.2  评价指标与实验局限", "Heading 3"),
    ("程序同时输出训练集和测试集预测准确率，并绘制分类结果图。准确率能够快速衡量分类器整体性能，但对于小样本数据集而言，单次划分可能导致评估结果受随机种子影响。更严谨的实验可采用交叉验证报告平均准确率与方差，并进一步计算混淆矩阵，以观察模型是否主要在相邻花种之间发生误判。", "Normal"),
]
insert_items_before(anchors["1.2  Session02"], session1)

session2 = [
    ("1.2.1  回归目标与损失函数分析", "Heading 3"),
    ("Boston Housing 实验从分类问题转向连续值预测。模型目标不是输出离散类别，而是估计房价数值，因此损失函数选择均方误差。均方误差对大偏差样本具有二次惩罚效果，能够迫使模型关注预测偏离明显的样本；但其缺点是对异常点敏感，若数据中存在极端房价或特征噪声，训练曲线可能出现较大波动。", "Normal"),
    ("线性回归层本质上学习特征到房价的加权组合。与深层网络相比，该模型表达能力有限，却具有参数含义清晰、训练稳定和计算开销低的优点。对于入门实验而言，它能够将数据读取、前向传播、损失计算、反向传播和参数保存串联成完整监督学习流程。", "Normal"),
    ("1.2.2  静态图到动态图的实现演进", "Heading 3"),
    ("三个 Session02 脚本体现了 PaddlePaddle 接口范式的演进。早期 fluid 静态图实现需要先定义计算图，再由 Executor 执行训练过程；动态图版本则通过继承 paddle.nn.Layer 构建 LinearRegression 类，在 Python 控制流中直接执行前向计算和反向传播。后者更符合当前深度学习框架的主流使用方式，也更便于调试单步损失、查看张量形状和保存模型参数。", "Normal"),
    ("训练损失曲线和预测散点图分别对应模型优化过程与泛化表现。若损失曲线整体下降，说明优化器能够沿着梯度方向降低误差；若预测点大致分布在真实值对角线附近，说明模型学习到了主要线性关系。图中仍存在离散误差，表明单层线性模型难以完全刻画房价与多维社会经济特征之间的非线性关系。", "Normal"),
]
insert_items_before(anchors["1.3  Session03"], session2)

session3 = [
    ("1.3.1  CNN 结构设计分析", "Heading 3"),
    ("CIFAR-10 图像尺寸小但类别差异复杂，单纯多层感知机难以有效利用局部空间结构。脚本采用三层卷积网络，先由浅层卷积提取边缘、颜色块和局部纹理，再通过更高通道数的卷积层组合局部模式。BatchNorm 的引入有助于稳定中间特征分布，MaxPool 则降低空间分辨率并扩大感受野。", "Normal"),
    ("全连接层输入 50 × 4 × 4 的特征张量并输出 10 类 logits，说明前序卷积模块已将图像压缩为较小的语义表征。损失函数采用交叉熵，符合多类互斥分类任务；优化器采用 Adam，可自适应调整参数更新尺度，通常比普通 SGD 在初期训练中更快降低损失。", "Normal"),
    ("1.3.2  训练稳定性与可复现实验条件", "Heading 3"),
    ("该实验设置 batch size = 128、epoch = 20、learning rate = 0.001，属于兼顾训练速度与稳定性的配置。由于 CIFAR-10 相比 MNIST 噪声更高、类间差异更细，训练准确率和测试准确率之间的差距需要重点观察。若训练准确率持续上升而测试准确率停滞，则说明模型可能过拟合；后续可加入数据增强、Dropout、权重衰减或更深的残差结构。", "Normal"),
    ("脚本中的推理阶段尝试读取指定图片并输出类别，这使训练结果从批量测试扩展到单样本应用。该阶段也暴露出工程复现要求：测试图片路径、图像归一化方式和训练阶段预处理必须一致，否则模型输出会受到输入分布偏移影响。", "Normal"),
]
insert_items_before(anchors["1.4  Session04"], session3)

session4 = [
    ("1.4.1  序列建模机制分析", "Heading 3"),
    ("IMDB 情感分类实验处理的是词序列而非定长数值表。脚本先利用词典将文本映射为整数 ID，再由网络将 ID 序列转换为隐状态表示。LSTM 的核心作用是通过门控结构缓解普通循环网络的长期依赖衰减问题，使模型能够综合评论前后文信息判断情感倾向。", "Normal"),
    ("与图像卷积网络强调局部空间模式不同，LSTM 强调时间或词序上的上下文传递。评论中某些情感词可能受到否定词、转折词或长距离修饰影响，因此仅统计词频不足以完成可靠判断。dynamic_lstm 能够处理变长序列，使模型不必将所有评论强行截断为完全相同的语义长度。", "Normal"),
    ("1.4.2  推理输出与误差来源", "Heading 3"),
    ("程序在预测阶段输出正面概率与负面概率，这比单一标签更有解释价值。若两个概率接近，说明样本位于模型判别边界附近，可能包含混合情绪或词典无法覆盖的表达。误差来源主要包括词典外词汇映射为 <unk>、训练语料与自定义评论存在语言风格差异，以及静态图接口对实验调试不如动态图直观。", "Normal"),
]
insert_items_before(anchors["1.5  Session05"], session4)

session5 = [
    ("1.5.1  基线模型意义", "Heading 3"),
    ("MNIST 实验采用两层隐层的多层感知机，而未使用卷积结构。该设计的价值在于建立图像分类的全连接基线：模型先将二维像素展开为一维向量，再通过非线性隐层学习数字形状与类别之间的映射。虽然展平操作破坏了像素邻域结构，但 MNIST 图像背景简单、数字居中，因此 MLP 仍能取得可接受效果。", "Normal"),
    ("与 Session03 的 CNN 相比，MNIST MLP 更适合说明高层 Model API 的训练流程。paddle.Model 封装了训练、评估和指标统计，使实验者可以集中理解网络结构、损失函数和评价指标之间的关系。该封装降低了代码复杂度，但也减少了对每一步梯度更新细节的显式控制。", "Normal"),
    ("1.5.2  数据预处理与泛化分析", "Heading 3"),
    ("脚本使用 Normalize 对图像进行标准化，使输入分布更适合神经网络训练。若忽略归一化，像素尺度差异可能导致梯度更新不稳定。模型在测试样本上的单例预测展示了从训练模型到实际推断的闭环，但单个样本不能代表整体泛化性能。若需要形成严格实验结论，应报告完整测试集准确率、错误样本示例及不同随机初始化下的稳定性。", "Normal"),
]
insert_items_before(anchors["2  最终四选二"], session5)

mission2_more = [
    ("2.1.4  状态空间、奖励函数与策略质量分析", "Heading 3"),
    ("Mission02 的状态空间由网格坐标定义，动作空间固定为四个方向，因而适合使用表格型 Q-Learning。字典式 Q 表只为访问过的状态分配价值向量，在障碍物较多或迷宫规模变化时比完整三维数组更节省空间。该结构也使状态键可以直接扩展为更复杂的元组，例如加入动态障碍物位置或智能体历史状态。", "Normal"),
    ("奖励函数决定智能体学习目标的形状。到达终点奖励 +100 提供稀疏但强烈的成功信号；障碍物惩罚 -10 使无效动作价值降低；普通移动 -1 将路径长度转化为累计回报差异，从而鼓励短路径。若去掉步长惩罚，智能体只需到达终点即可，可能学习到明显绕行策略；若障碍物惩罚过大，智能体在早期探索中可能过度保守。当前奖励设计在路径效率和探索容错之间取得折中。", "Normal"),
    ("2.1.5  图形化界面与运行结果预留", "Heading 3"),
]
insert_items_before(anchors["2.1.4"], mission2_more)
insert_placeholder_before(
    doc,
    anchors["2.1.4"],
    "图 2-1  Mission02 Q-Learning 迷宫寻路系统最终图形化界面（待插入）",
    "为呈现最终运行效果，Mission02 在算法分析之后预留图形化界面截图位置。该图建议展示迷宫网格、起点、终点、障碍物、训练控制面板、回合计数以及收敛后绘制的最优路径。",
)
anchors["2.1.4"].text = "2.1.6  工程权衡与边界条件讨论"

mission3_more = [
    ("2.2.4  知识库完备性与解释机制分析", "Heading 3"),
    ("Mission03 的知识库采用分层产生式结构，基础特征首先推出哺乳动物、鸟类、食肉动物和有蹄类动物等中间概念，再由中间概念与外观特征共同推出目标动物。这种结构减少了规则重复：例如多个哺乳类动物可以共享“有毛发”或“产奶”推出哺乳动物的基础规则，而不必在每条目标规则中重复全部原始特征。", "Normal"),
    ("解释日志是专家系统区别于黑箱分类模型的重要特征。系统不仅给出最终动物名称，还记录初始事实、触发规则、加入事实和终止条件，使推理链可以被人工审查。当证据不足时，候选推荐机制通过目标依赖集合与当前事实库的交集提供部分匹配结果，避免将不完整输入简单判定为失败。该设计提升了系统在教学演示中的可解释性。", "Normal"),
    ("2.2.5  图形化界面与运行结果预留", "Heading 3"),
]
insert_items_before(anchors["2.2.4"], mission3_more)
insert_placeholder_before(
    doc,
    anchors["2.2.4"],
    "图 2-2  Mission03 动物识别专家系统最终图形化界面（待插入）",
    "为呈现最终运行效果，Mission03 在推理机制分析之后预留图形化界面截图位置。该图建议展示特征选择区、推理按钮、解释日志、候选推荐信息和最终识别结果。",
)
anchors["2.2.4"].text = "2.2.6  推理策略与扩展文件说明"

# Remove appendix entry from manual TOC if still present.
for p in list(doc.paragraphs):
    if p.text.strip() == "附录A  源文件覆盖清单":
        remove_element(p._element)

doc.save(str(DOCX))
print(DOCX)
