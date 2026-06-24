"""
Generate 改错题.docx — comprehensive summary of ALL correction questions
across 样卷1-8, 提升卷1-3, 提高卷4 (12 papers total).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_DIR = 'D:/ScuderiaAMG/java/java'


def set_run_font(run, font_name='等线', font_size=9, bold=False, color=None):
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_p(doc, text, font_size=9, bold=False, color=None, alignment=None,
          space_after=4, space_before=0, font_name='等线', indent=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    if indent:
        pf.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold, color)
    return p


def add_code(doc, text, font_size=7.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.1
    pf.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Consolas'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Consolas')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    return p


def add_section(doc, text, font_size=11):
    return add_p(doc, text, font_size=font_size, bold=True, space_before=10, space_after=4)


def add_subsection(doc, text, font_size=9):
    return add_p(doc, text, font_size=font_size, bold=True, space_before=6, space_after=3)


# ========================================================================
# Build document
# ========================================================================
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Title
add_p(doc, 'Java程序改错题 分类总结（完整版）', font_size=15, bold=True,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_p(doc, '覆盖样卷1-8、提升卷1-3、提高卷4 | 共12套试卷 · 24道改错条目 | 按错误类型分13类',
      font_size=7.5, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
      color=(100, 100, 100))

# ========================================================================
# 一、访问权限错误
# ========================================================================
add_section(doc, '一、访问权限错误（private成员外部访问）', font_size=11)
add_subsection(doc, '▎来源：样卷1第1题 · 样卷6第1题')
add_p(doc, '【错误特征】在类的外部直接访问另一个类的 private 成员变量（如 obj.field = value），编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】private 限定成员仅在本类内部可见，编译器在语义分析阶段检查访问权限。外部类即使是同包也无法直接读写 private 成员。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】删除直接访问语句，改为调用 public setter；或为类添加相应的 public 方法。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：直接访问private成员
Calculator c = new Calculator();
c.result = 100;              // 编译错误

// 修改：添加setter方法
public void setResult(int r) { this.result = r; }
c.setResult(100);             // 正确''')

# ========================================================================
# 二、继承构造方法缺失
# ========================================================================
add_section(doc, '二、继承中的构造方法缺失（super调用失败）', font_size=11)
add_subsection(doc, '▎来源：样卷1第2题 · 样卷6第2题')
add_p(doc, '【错误特征】子类无显式构造方法，父类只有带参构造——编译错误：隐式 super() 找不到父类无参构造。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】子类构造第一条必须是 super(...) 或 this(...)。无显式构造时编译器生成默认构造，隐式调用 super()。若父类定义了带参构造而未定义无参构造（此时编译器不会自动生成），super() 就找不到匹配的父类构造。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】子类显式定义构造方法，第一行 super(参数) 调用父类已有的带参构造。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：父类只有带参构造，子类无显式构造
class Person {
    public Person(String name) { ... }   // 无无参构造
}
class Student extends Person {
    int grade;                           // 编译错误
}

// 修改：显式构造 + super(name)
class Student extends Person {
    int grade;
    public Student(String name) {
        super(name);                     // 显式调用父类构造
    }
}''')

# ========================================================================
# 三、抽象类/接口被实例化
# ========================================================================
add_section(doc, '三、抽象类或接口被直接实例化', font_size=11)
add_subsection(doc, '▎来源：样卷2第1题（抽象类）· 样卷2第2题（接口）')
add_p(doc, '【错误特征】使用 new 直接创建抽象类或接口的对象（如 new AbstractClass() 或 new InterfaceName()），编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】抽象类包含未实现的方法（abstract method），接口中所有方法默认都是抽象的——它们都是"不完整的类型"。Java禁止直接实例化不完整的类型，因为如果允许 new，调用其抽象方法时会没有代码可执行。只有具体子类/实现类才能被 new。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】改为创建具体子类或实现类的对象：AbstractClass a = new ConcreteSubclass(); 或 Interface i = new ImplClass();',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：抽象类和接口不能new
Vehicle v = new Vehicle();          // Vehicle是abstract class → 错误
Calculator c = new Calculator();    // Calculator是interface → 错误

// 修改：创建子类/实现类对象
Vehicle v = new Car();              // Car extends Vehicle
Calculator c = new SimpleCalculator(); // SimpleCalculator implements Calculator''')

# ========================================================================
# 四、接口方法未全部实现
# ========================================================================
add_section(doc, '四、接口方法未全部实现', font_size=11)
add_subsection(doc, '▎来源：样卷2第2题 · 样卷7第1题')
add_p(doc, '【错误特征】类 implements 接口但未实现全部抽象方法，且该类不是抽象类——编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】实现接口=签署契约。具体类必须为接口每个方法提供 public 实现（方法签名完全一致），缺一不可。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】补全所有缺失方法，访问权限必须是 public（接口方法默认 public abstract）。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：缺少接口方法实现
interface Calculator {
    int add(int a, int b);
    int subtract(int a, int b);
}
class SimpleCalculator implements Calculator {
    public int add(int a, int b) { return a + b; }
    // 缺少subtract() → 编译错误
}

// 修改：补全所有方法（必须public）
class SimpleCalculator implements Calculator {
    public int add(int a, int b) { return a + b; }
    public int subtract(int a, int b) { return a - b; }
}''')

# ========================================================================
# 五、父类引用无法访问子类特有方法
# ========================================================================
add_section(doc, '五、向上转型后无法调用子类特有方法', font_size=11)
add_subsection(doc, '▎来源：样卷1第2题')
add_p(doc, '【错误特征】Animal a = new Dog(); 后调用 a.wagTail()——wagTail() 只在 Dog 中定义，Animal 中没有，编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】编译时，编译器只看引用类型（Animal），不看实际对象类型（Dog）。Animal 类中没有 wagTail() 方法，因此编译失败——这是静态类型检查的约束。方法调用=编译时检查引用类型是否有该方法，运行时再动态绑定到实际对象的实现。但前提是引用类型必须有该方法的声明。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】①向下转型：((Dog) a).wagTail()（需确保a实际是Dog）；②改用子类引用：Dog a = new Dog("旺财")。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：父类引用无法调用子类特有方法
Animal a = new Dog();
a.wagTail();                   // 编译错误：Animal没有wagTail()

// 修改方式一：向下转型
((Dog) a).wagTail();           // 正确（需确保a实际为Dog）

// 修改方式二：使用子类引用
Dog a = new Dog("旺财");
a.wagTail();                   // 正确''')

# ========================================================================
# 六、静态方法隐藏
# ========================================================================
add_section(doc, '六、静态方法隐藏（误以为可重写）', font_size=11)
add_subsection(doc, '▎来源：样卷7第2题')
add_p(doc, '【错误特征】子类定义与父类同名的 static 方法，期望父类引用调用时触发多态——实际调用的是父类版本（编译时绑定）。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】static 方法属于类而非实例，调用在编译时根据引用类型决定（隐藏 hide），而非运行时动态绑定（重写 override）。Parent p = new Child(); p.staticMethod() → 编译时 p 是 Parent 类型，直接绑定 Parent 的静态方法。@Override 对静态方法报错。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】去掉 static 改为实例方法，子类同名方法即真正重写，运行时多态生效。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：静态方法隐藏，非多态
Parent p = new Child();
p.info();   // 输出 "Parent info"（编译时绑定），不是期望的"Child info"

// 修改：去掉static → 实例方法重写
class Parent {
    public void info() { ... }     // 非static
}
class Child extends Parent {
    @Override public void info() { ... }
}
p.info();   // 输出 "Child info"（多态生效）''')

# ========================================================================
# 七、已检查异常未处理
# ========================================================================
add_section(doc, '七、已检查异常（Checked Exception）未处理', font_size=11)
add_subsection(doc, '▎来源：样卷4第1题 · 样卷5第1题 · 样卷8第1题 · 提高卷4第1题')
add_p(doc, '【错误特征】调用可能抛出 IOException / SQLException / InterruptedException / ClassNotFoundException 等的方法，既无 try-catch 也无 throws 声明，编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】Java "Catch or Specify" 原则：Exception 及子类（不含 RuntimeException）为已检查异常，必须编译时处理——try-catch 当场捕获，或 throws 声明向上传播。编译器强制检查，违反则不生成 .class。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】用 try-catch 包围，或在方法签名加 throws 声明（最终调用处推荐 try-catch）。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：Checked异常未处理，编译错误
public static void main(String[] args) {
    Thread.sleep(5000);                  // InterruptedException
    FileWriter fw = new FileWriter("");  // IOException
}

// 修改：try-catch捕获
public static void main(String[] args) {
    try {
        Thread.sleep(5000);
        FileWriter fw = new FileWriter("out.txt");
        fw.write("data"); fw.close();
    } catch (InterruptedException | IOException e) {
        e.printStackTrace();
    }
}''')

# ========================================================================
# 八、I/O与JDBC资源未关闭
# ========================================================================
add_section(doc, '八、I/O与JDBC资源未关闭（资源泄漏）', font_size=11)
add_subsection(doc, '▎来源：样卷4第2题 · 样卷8第1题 · 提高卷4第1题 · 提高卷4第2题')
add_p(doc, '【错误特征】使用 FileWriter / Stream / Connection / Statement / ResultSet / ObjectOutputStream 后不调用 close()；或缺少 flush()。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】这些资源占用OS句柄、端口、缓冲区。不关闭→①缓冲数据未刷新丢失；②句柄泄漏→"Too many open files"；③连接池耗尽。JDBC多层资源需按"先开后关"顺序（ResultSet→Statement→Connection）。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】推荐 try-with-resources（Java 7+），或 try-catch-finally 中判 null 后 close()。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：多层资源未关，驱动未加载
Connection conn = DriverManager.getConnection(url, user, pass);
Statement stmt = conn.createStatement();
ResultSet rs = stmt.executeQuery("SELECT * FROM t");
// 全都没关！

// 修改：try-with-resources（推荐）
try (Connection conn = DriverManager.getConnection(url, user, pass);
     Statement stmt = conn.createStatement();
     ResultSet rs = stmt.executeQuery("SELECT * FROM t")) {
    while (rs.next()) { ... }
} catch (SQLException e) { e.printStackTrace(); }
// 自动按ResultSet→Statement→Connection顺序关闭''')

# ========================================================================
# 九、泛型缺失与类型不安全
# ========================================================================
add_section(doc, '九、泛型缺失导致类型不安全（Raw Type → ClassCastException）', font_size=11)
add_subsection(doc, '▎来源：样卷3第1题 · 样卷8第2题 · 提升卷2第2题')
add_p(doc, '【错误特征】使用原始类型 List/Map（无 < >），混入不同类型元素，遍历时强制转换 → 运行时 ClassCastException。泛型集合中 add 不匹配类型 → 编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】无泛型时集合元素类型退化为 Object，编译器不检查 add 类型——Integer 和 String 都能放入 List。取出时强制转换，若实际类型不兼容则抛 ClassCastException。泛型将类型检查提前到编译期（类型擦除后运行时仍为 Object，但编译期检查已保障安全）。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】声明泛型参数 List<String>，编译器阻止不匹配的 add，取出时无需强制转换。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误一：原始类型→运行时期ClassCastException
List list = new ArrayList();
list.add("Hello");
list.add(123);               // Integer混入
String s = (String) list.get(1);  // ClassCastException!

// 修改：加泛型
List<String> list = new ArrayList<>();
list.add("Hello");
// list.add(123);           // 编译错误，提前发现！

// 错误二：泛型集合add不匹配类型
List<String> strList = new ArrayList<>();
strList.add(200);            // 编译错误，int不能放List<String>

// 修改：改为String
strList.add("200");          // 正确''')

# ========================================================================
# 十、自动拆箱 null → NullPointerException
# ========================================================================
add_section(doc, '十、自动拆箱 null 引发 NullPointerException', font_size=11)
add_subsection(doc, '▎来源：样卷3第2题')
add_p(doc, '【错误特征】HashMap 中 remove 后使用 int 接收 get() 返回值，键已删除 get() 返回 null，自动拆箱 null → int 时抛出 NullPointerException。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】Java 自动拆箱（Auto-unboxing）将包装类型（Integer）自动转为基本类型（int）时，实际调用 intValue() 方法。如果包装对象为 null，调用 intValue() 会抛出 NullPointerException。常见场景：Map.get() 对不存在的键返回 null 后直接赋给 int，或方法返回 Integer 直接用 int 接收。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】用包装类型 Integer 接收可能为 null 的返回值，使用前进行 null 检查。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：null自动拆箱→NullPointerException
map.remove("Li");
int deleted = map.get("Li");    // get返回null，拆箱→NPE!

// 修改：用包装类型 + null检查
Integer deleted = map.get("Li");
if (deleted != null) {
    System.out.println("After remove: " + deleted);
} else {
    System.out.println("Li 已被删除，不存在");
}''')

# ========================================================================
# 十一、序列化问题
# ========================================================================
add_section(doc, '十一、对象序列化相关问题', font_size=11)
add_subsection(doc, '▎来源：样卷5第2题 · 提升卷2第1题 · 提高卷4第2题')
add_p(doc, '【错误特征】①未实现 Serializable 的类被 writeObject() 写入→运行时 NotSerializableException；②无 serialVersionUID →版本兼容风险；③ObjectStream 未关闭；④未理解 transient 含义。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】①ObjectOutputStream.writeObject() 要求对象实现 Serializable 标记接口（无方法，仅标记"可序列化"），否则抛 NotSerializableException。String、Integer、ArrayList 等JDK类均已实现。②无显式 serialVersionUID 时编译器自动计算hash，类改动后hash变化→旧文件反序列化抛 InvalidClassException。③transient 字段跳过序列化，反序列化后为默认值（null/0/false），适用于密码/缓存。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】类 implements Serializable + 显式定义 serialVersionUID + 敏感字段加 transient + try-with-resources 关闭流。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：未实现Serializable → NotSerializableException
class Student {
    String name; int age;
}
oos.writeObject(new Student());    // 运行时异常！

// 修改：实现Serializable + 加UID + transient
class Student implements Serializable {
    private static final long serialVersionUID = 1L;
    private String name;
    private transient int age;       // 不参与序列化
    // ...
}
try (ObjectOutputStream oos = new ObjectOutputStream(
         new FileOutputStream("s.ser"))) {
    oos.writeObject(new Student("Tom", 20));
    oos.flush();
} catch (IOException e) { e.printStackTrace(); }
// 反序列化后 age = 0（transient默认值）''')

# ========================================================================
# 十二、protected 跨包访问规则
# ========================================================================
add_section(doc, '十二、protected 跨包访问限制', font_size=11)
add_subsection(doc, '▎来源：提升卷3第1题')
add_p(doc, '【错误特征】跨包子类中通过父类引用（new 父类()）访问 protected 方法→编译错误。用 this 访问则合法。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】protected 的精细规则：同包子类可访问父类对象的 protected 成员。但不同包子类只能通过自身引用（或子类引用）访问继承来的 protected 成员，不能通过父类引用去访问另一个父类对象的 protected 成员。这是 Java 为防止跨包绕过访问控制的安全设计——如果允许通过父类引用访问，任意代码都可以 new 一个父类对象然后访问其 protected 成员。',
      font_size=8, indent=0.3)
add_p(doc, '【如何修改】改为 this.method() 或直接 method()，通过自身引用访问继承的 protected 成员。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：跨包子类用父类引用访问protected成员
// AppMain extends BaseLib（不同包）
public void run() {
    BaseLib base = new BaseLib();
    base.helper();          // 编译错误！用父类引用访问
}

// 修改：通过自身引用（或直接调用）
public void run() {
    this.helper();          // 合法：通过继承访问
    // 或直接 helper();
}''')

# ========================================================================
# 十三、接口默认方法冲突 & 异常重写规则（知识点辨析）
# ========================================================================
add_section(doc, '十三、接口默认方法冲突 & 异常重写规则（辨析类）', font_size=11)
add_subsection(doc, '▎来源：提升卷1第1题 · 提升卷1第2题 · 提升卷3第2题')

add_p(doc, '【子类A — 多接口默认方法冲突】', font_size=8.5, indent=0.3, bold=True, space_before=4)
add_p(doc, '【错误特征】一个类同时实现两个接口，两接口有同名 default 方法——编译错误（编译器无法决定用哪个）。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】Java 不允许多继承的"菱形问题"在默认方法上重演。当多个接口提供同名默认实现时，类必须显式重写该方法以消除歧义。可在重写中通过 InterfaceName.super.method() 选择某个接口的版本。',
      font_size=8, indent=0.3)
add_p(doc, '【修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// 错误：两个接口有同名default方法→编译冲突
interface A { default void info() { ... } }
interface B { default void info() { ... } }
class C implements A, B { }   // 编译错误！

// 修改：必须重写info()，可指定用哪个接口的版本
class C implements A, B {
    @Override
    public void info() {
        A.super.info();       // 选择A的版本
        // 或自己实现
    }
}''')

add_p(doc, '【子类B — 类优先原则（辨析）】', font_size=8.5, indent=0.3, bold=True, space_before=6)
add_p(doc, '【场景】类继承父类（已有方法实现），同时实现含同名默认方法的接口。',
      font_size=8, indent=0.3)
add_p(doc, '【原理】类优先原则（Class wins）：父类中的实现 > 接口中的默认方法。编译器选择父类版本，不报错。这是 Java 解决"菱形继承"问题的核心规则。',
      font_size=8, indent=0.3)

add_p(doc, '【子类C — 异常重写规则（辨析）】', font_size=8.5, indent=0.3, bold=True, space_before=6)
add_p(doc, '【错误特征】子类重写方法声明的异常比父类更宽泛→编译错误。',
      font_size=8, indent=0.3)
add_p(doc, '【错误原理】重写规则对异常的限制：子类方法不能抛出比父类方法更宽泛的已检查异常（可抛出更窄的或不抛出，RuntimeException 不受限）。这是为了多态安全——调用者按父类声明处理异常，子类突然抛更宽异常会绕过catch。',
      font_size=8, indent=0.3)
add_p(doc, '【三个辨析修改示例】', font_size=8, indent=0.3, bold=True)
add_code(doc, '''// === 默认方法冲突 ===（提升卷1-2.2）
// 错误：Diagram implements Drawable,Scalable → 两接口都有printInfo()
// 修改：显式重写
@Override public void printInfo() {
    Drawable.super.printInfo();   // 或自定义
}

// === 类优先原则 ===（提升卷3-2.2，此题无实际错误）
class Middle implements A { @Override public void common() {...} }
class Final extends Middle implements B { }
// Final从Middle继承common()，接口B的common()被忽略——类优先

// === 异常重写规则 ===（提升卷1-2.1）
// 错误：父类throws Exception，子类throws Throwable（更宽）
// 修改：子类throws Exception或更窄（或不抛）
@Override
public void process() throws Exception {    // 与父类一致
    ...
}''')

# ========================================================================
# 附录：完整速查表
# ========================================================================
add_section(doc, '附：全部改错条目速查表（按试卷排列）', font_size=11)

# 24 rows + 1 header = 25 rows
table = doc.add_table(rows=26, cols=5)
table.style = 'Table Grid'

headers = ['来源', '错误类型', '核心特征', '原理关键词', '修改方式']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_run_font(run, '等线', 7, bold=True)

data = [
    # 样卷1
    ['样卷1-1', '构造缺失+private', '无参构造缺失；直接访问private字段', '默认构造+super()；private类内可见', '加无参构造或传参；加setter'],
    ['样卷1-2', '继承+向上转型', '子类缺super(name)；父引用调子类方法', '隐式super()找无参；编译时按引用类型', '显式super(name)；向下转型或改引用'],
    # 样卷2
    ['样卷2-1', '抽象类实例化', 'new Vehicle()—Vehicle是abstract', '抽象类不完整，不可直接实例化', 'new Car()—子类对象赋父类引用'],
    ['样卷2-2', '接口未实现+实例化', '只实现了部分方法；new接口', '接口=全部抽象方法；接口不可实例化', '补全所有方法；new实现类'],
    # 样卷3
    ['样卷3-1', '原始类型', 'List无泛型→Integer/String混放→强转崩溃', '绕过编译检查→ClassCastException', 'List<Integer>，编译即阻止错误add'],
    ['样卷3-2', 'null自动拆箱', 'map.remove后get返回null→int接收→NPE', '自动拆箱调用intValue()，null.xxx→NPE', '用Integer接收+null检查'],
    # 样卷4
    ['样卷4-1', 'Checked异常', 'FileWriter抛出IOException未处理', 'Catch or Specify编译期强制', 'try-catch-finally或throws声明'],
    ['样卷4-2', 'JDBC异常+资源', 'SQLException未处理+无驱动+未关资源', '多层异常+多层资源先后顺序', 'try-catch-finally+依次close+加载驱动'],
    # 样卷5
    ['样卷5-1', '中断异常', 'Thread.sleep()抛InterruptedException未处理', 'sleep()声明了throws，必须处理', 'try { Thread.sleep(); } catch(...) {}'],
    ['样卷5-2', '序列化', 'new Object()未实现Serializable→写入失败', 'writeObject要求Serializable标记接口', '只写可序列化对象(如Integer/String)'],
    # 样卷6
    ['样卷6-1', '访问权限', '外部直接设c.result=100，result是private', 'private仅在类内可见', '加public setResult()方法'],
    ['样卷6-2', '继承构造', 'Student extends Person，Person只有带参构造', '父类无无参构造+子类缺显式构造→编译错误', '在Student加构造+super(name)'],
    # 样卷7
    ['样卷7-1', '接口实现', 'Circle implements Drawable但缺draw()', '具体类必须实现接口全部方法', '补全public void draw(){...}'],
    ['样卷7-2', '静态隐藏', 'p.info()调了父类版，因子类static是隐藏非重写', 'static方法编译时绑定引用类型', '去static→实例方法多态'],
    # 样卷8
    ['样卷8-1a', 'Checked异常', 'FileWriter构造抛IOException未处理', '已检查异常必须catch或throws', 'try-catch或throws IOException'],
    ['样卷8-1b', '资源泄漏', 'FileWriter用后未close()', '句柄泄漏+数据可能未刷新', 'try-with-resources / finally close'],
    ['样卷8-2', '原始类型→CCE', 'raw List混放String/Integer→强转String崩', 'Integer cannot be cast to String', 'List<String>，编译期类型安全'],
    # 提高卷4
    ['提高卷4-1', 'JDBC全套', '无try-catch+无驱动加载+未关三层资源', 'SQLException+ClassNotFoundException+泄漏', 'try-catch-finally+驱动+依次close'],
    ['提高卷4-2', '序列化全套', '无serialVersionUID+流未关+无flush', '版本不兼容+数据不完整+资源泄漏', '定义UID+try-with-resources+flush'],
    # 提升卷1
    ['提升卷1-1', '异常重写规则', '子类throws比父类更宽的已检查异常', '重写异常不能更宽(RuntimeException除外)', '子类throws与父类一致或更窄'],
    ['提升卷1-2', '默认方法冲突', '两接口同名default方法→编译器无法选择', '多继承default方法歧义', '类中必须显式重写，可用X.super.m()'],
    # 提升卷2
    ['提升卷2-1', '序列化', 'Student未implements Serializable写文件', 'writeObject→NotSerializableException', 'implements Serializable+定义UID'],
    ['提升卷2-2', '泛型+原始类型', 'raw List混三类对象；泛型List加int', '编译警告→CCE；编译错误直接阻止', '统一泛型+类型匹配的add参数'],
    # 提升卷3
    ['提升卷3-1', 'protected跨包', '跨包子类用父类引用调protected方法', '跨包只能通过自身引用访问继承的protected', '改用this.method()或直接调用'],
    ['提升卷3-2', '类优先(辨析)', '父类已实现，接口默认方法被忽略，无错误', 'Class wins原则→类方法优先于接口默认', '辨析知识点：不会报错，正常编译运行'],
]

for r, row_data in enumerate(data):
    for c, text in enumerate(row_data):
        cell = table.rows[r + 1].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_run_font(run, '等线', 6.5, bold=False)

# ========================================================================
# 错误类型分布统计
# ========================================================================
add_section(doc, '附二：错误类型分布统计', font_size=10)

stats_data = [
    ['Checked异常未处理', '样卷4-1, 样卷5-1, 样卷8-1a, 提高卷4-1', '4次'],
    ['资源未关闭（I/O+JDBC）', '样卷4-2, 样卷8-1b, 提高卷4-1, 提高卷4-2', '4次'],
    ['泛型缺失/原始类型', '样卷3-1, 样卷8-2, 提升卷2-2', '3次'],
    ['访问权限（private/protected）', '样卷1-1, 样卷6-1, 提升卷3-1', '3次'],
    ['序列化问题', '样卷5-2, 提升卷2-1, 提高卷4-2', '3次'],
    ['继承构造super()缺失', '样卷1-2, 样卷6-2', '2次'],
    ['接口方法未全部实现', '样卷2-2, 样卷7-1', '2次'],
    ['抽象类/接口被实例化', '样卷2-1, 样卷2-2', '2次'],
    ['静态方法隐藏', '样卷7-2', '1次'],
    ['向上转型限制', '样卷1-2', '1次'],
    ['null自动拆箱→NPE', '样卷3-2', '1次'],
    ['接口默认方法冲突', '提升卷1-2', '1次'],
    ['异常重写规则', '提升卷1-1', '1次'],
    ['类优先原则(辨析)', '提升卷3-2', '1次'],
]

stats_table = doc.add_table(rows=len(stats_data)+1, cols=3)
stats_table.style = 'Table Grid'

for i, h in enumerate(['错误类型', '出现位置', '频次']):
    cell = stats_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_run_font(run, '等线', 7.5, bold=True)

for r, row_data in enumerate(stats_data):
    for c, text in enumerate(row_data):
        cell = stats_table.rows[r + 1].cells[c]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_run_font(run, '等线', 7, bold=False)

# ========================================================================
# Save
# ========================================================================
output_path = os.path.join(OUTPUT_DIR, '改错题.docx')
doc.save(output_path)
print(f'Created: {output_path}')
print(f'Total: 12 papers, 24 correction item entries, 13 error categories')
