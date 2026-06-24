"""
Generate exam papers: 样卷6, 样卷7, 样卷8, 提高卷4 and their answer files.
Each paper: 7 reading questions + 2 correction questions + 2 design questions.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_DIR = 'D:/ScuderiaAMG/java/java'


def set_cell_font(run, font_name='等线', font_size=10.5, bold=False):
    """Set font for a run."""
    run.font.size = Pt(font_size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph(doc, text, font_name='等线', font_size=10.5, bold=False, alignment=None, space_after=6, space_before=0):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.25
    run = p.add_run(text)
    set_cell_font(run, font_name, font_size, bold)
    return p


def add_code_block(doc, code_text, font_size=9):
    """Add a code block paragraph with monospace feel."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)
    pf.line_spacing = 1.15
    pf.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.size = Pt(font_size)
    run.font.name = 'Consolas'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Consolas')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    return p


def add_section_title(doc, text):
    """Add a section title (一、二、三)."""
    return add_paragraph(doc, text, font_size=12, bold=True, space_before=12, space_after=6)


def add_question_title(doc, text):
    """Add a question title."""
    return add_paragraph(doc, text, font_size=10.5, bold=True, space_before=8, space_after=4)


def add_body_text(doc, text):
    """Add normal body text."""
    return add_paragraph(doc, text, font_size=10.5, space_after=4)


def create_paper(paper_data, filename):
    """Create an exam paper docx file."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '等线'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    add_paragraph(doc, paper_data['title'], font_size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, '总分：100分 | 考试时间：120分钟', font_size=10.5,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # === Section 1: Reading Questions ===
    add_section_title(doc, '一、看程序写结果（共42分）')
    add_body_text(doc, '阅读下列程序，写出程序的运行结果。')

    for i, q in enumerate(paper_data['reading']):
        add_question_title(doc, f'{i+1}.（{q["score"]}分）阅读以下程序，写出运行结果。')
        add_code_block(doc, q['code'])
        if 'question' in q:
            add_body_text(doc, q['question'])
        add_paragraph(doc, '', font_size=6, space_after=6)  # spacer

    # === Section 2: Correction Questions ===
    add_section_title(doc, '二、程序改错（共20分）')
    add_body_text(doc, '找出下列程序中的错误，说明错误原因并给出修改办法。')

    for i, q in enumerate(paper_data['correction']):
        add_question_title(doc, f'{i+1}.（{q["score"]}分）')
        add_code_block(doc, q['code'])
        add_paragraph(doc, '', font_size=6, space_after=6)

    # === Section 3: Design Questions ===
    add_section_title(doc, '三、程序设计（共38分）')
    add_body_text(doc, '请根据要求编写完整的Java程序。')

    for i, q in enumerate(paper_data['design']):
        add_question_title(doc, f'{i+1}.（{q["score"]}分）{q["title"]}')
        add_body_text(doc, q['requirement'])
        add_paragraph(doc, '', font_size=6, space_after=6)

    doc.save(os.path.join(OUTPUT_DIR, filename))
    print(f'Created: {filename}')


def create_answer(paper_data, filename):
    """Create an answer file docx."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '等线'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_paragraph(doc, paper_data['title'] + ' 参考答案', font_size=16, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Reading answers
    add_section_title(doc, '一、看程序写结果（共42分）')
    for i, q in enumerate(paper_data['reading']):
        add_question_title(doc, f'{i+1}.（{q["score"]}分）运行结果')
        add_code_block(doc, q['answer'])
        if 'analysis' in q:
            add_body_text(doc, '解析：' + q['analysis'])
        add_paragraph(doc, '', font_size=6, space_after=4)

    # Correction answers
    add_section_title(doc, '二、程序改错（共20分）')
    for i, q in enumerate(paper_data['correction']):
        add_question_title(doc, f'{i+1}.（{q["score"]}分）')
        for err in q['errors']:
            add_body_text(doc, f'错误：{err["desc"]}')
            add_body_text(doc, f'修改办法：{err["fix"]}')
        if 'corrected_code' in q:
            add_body_text(doc, '修正后的代码：')
            add_code_block(doc, q['corrected_code'])
        add_paragraph(doc, '', font_size=6, space_after=4)

    # Design answers
    add_section_title(doc, '三、程序设计（共38分）')
    for i, q in enumerate(paper_data['design']):
        add_question_title(doc, f'{i+1}.（{q["score"]}分）{q["title"]} — 参考代码')
        add_code_block(doc, q['reference_code'])
        add_paragraph(doc, '', font_size=6, space_after=4)

    doc.save(os.path.join(OUTPUT_DIR, filename))
    print(f'Created: {filename}')


# ============================================================
# PAPER 6: 样卷6 (Basic Level)
# ============================================================
PAPER6 = {
    'title': 'Java程序设计 样卷6',
    'reading': [
        {
            'score': '6',
            'code': '''public class ArraySum {
    public static void main(String[] args) {
        int[] a = {12, 5, 8, 15, 3, 9};
        int sum = 0;
        int max = a[0];
        for (int i = 0; i < a.length; i++) {
            sum += a[i];
            if (a[i] > max) {
                max = a[i];
            }
        }
        System.out.println("sum=" + sum);
        System.out.println("max=" + max);
        System.out.println("avg=" + (double)sum / a.length);
    }
}''',
            'answer': '''sum=52
max=15
avg=8.666666666666666''',
            'analysis': '从main方法开始运行，运行到int[] a = {12, 5, 8, 15, 3, 9}声明并初始化数组a，看向int sum = 0和int max = a[0]，这里将sum初始化为0，max初始化为a[0]=12。进入for循环，i从0开始遍历到a.length-1（即5），每次循环执行sum += a[i]（将a[i]的值累加到sum）和if (a[i] > max)判断。i=0时a[0]=12，sum变为12，max保持12（12不大于12）；i=1时a[1]=5，sum变为17，max保持12；i=2时a[2]=8，sum变为25；i=3时a[3]=15，sum变为40，max更新为15（因为15>12）；i=4时a[4]=3，sum变为43；i=5时a[5]=9，sum变为52。循环结束后，第一个输出语句System.out.println("sum=" + sum)输出sum=52；第二个输出max=15；第三个输出avg=52.0/6≈8.667（(double)sum将int转为double进行浮点除法）。'
        },
        {
            'score': '6',
            'code': '''public class ForPattern {
    public static void main(String[] args) {
        for (int i = 1; i <= 4; i++) {
            for (int j = 1; j <= i; j++) {
                System.out.print("* ");
            }
            System.out.println();
        }
        int total = 0;
        for (int k = 1; k <= 10; k++) {
            if (k % 3 == 0) {
                total += k;
            }
        }
        System.out.println("total=" + total);
    }
}''',
            'answer': '''*
* *
* * *
* * * *
total=18''',
            'analysis': '从main方法开始运行，首先进入第一个嵌套for循环，外层循环i从1到4。看向内层for循环（j从1到i），在这里面我们执行了System.out.print("* ")打印星号后跟空格，但不执行println换行（因为用的是print而非println），直到内层循环结束才执行外层循环体中的System.out.println()实现换行。因此i=1时打印1个"* "，i=2时打印2个"* "，依此类推，形成直角三角形图案。第一个嵌套循环结束后换行。接着看向第二个循环：int total = 0初始化累加变量，for循环k从1到10，在循环体中执行了if (k % 3 == 0)判断k是否能被3整除，但是不执行total += k（当k不是3的倍数时跳过）。k=3时total=3，k=6时total=9，k=9时total=18。最后System.out.println("total=" + total)输出total=18。'
        },
        {
            'score': '6',
            'code': '''public class ScoreGrade {
    public static void main(String[] args) {
        int score = 82;
        char grade;
        if (score >= 90) {
            grade = 'A';
        } else if (score >= 80) {
            grade = 'B';
        } else if (score >= 70) {
            grade = 'C';
        } else if (score >= 60) {
            grade = 'D';
        } else {
            grade = 'E';
        }
        System.out.println("Score: " + score + " -> Grade: " + grade);
    }
}''',
            'answer': 'Score: 82 -> Grade: B',
            'analysis': "从main方法开始运行，运行到int score = 82声明并初始化成绩变量，看向if-else if-else多分支结构，在这里面我们从上到下依次判断每个条件。首先判断if (score >= 90)，82>=90为false，但是不执行grade = 'A'分支；接着判断else if (score >= 80)，82>=80为true，所以执行grade = 'B'，执行完后跳过后面所有else if和else分支（这是if-else if-else的执行规则——一旦某个条件为true，执行对应分支后直接跳出整个结构）。因此后面的score>=70、score>=60和else分支都不会被执行。最后System.out.println输出Score: 82 -> Grade: B。这里的关键点是条件从严格到宽松排列（>=90→>=80→>=70→>=60→else），82先匹配到>=80这个条件。"
        },
        {
            'score': '6',
            'code': '''public class StringCompare {
    public static void main(String[] args) {
        String s1 = "Java";
        String s2 = "Java";
        String s3 = new String("Java");
        String s4 = "Ja" + "va";

        System.out.println(s1 == s2);
        System.out.println(s1 == s3);
        System.out.println(s1.equals(s3));
        System.out.println(s1 == s4);
        System.out.println(s1.length());
    }
}''',
            'answer': '''true
false
true
true
4''',
            'analysis': '从main方法开始运行，首先声明了4个String变量。看向s1 = "Java"和s2 = "Java"，在这里面由于它们是编译期确定的字符串字面量，JVM将它们放入字符串常量池（String Pool），并且s1和s2指向常量池中同一个"Java"对象。s3 = new String("Java")则是在堆上创建了一个全新的String对象，虽然内容也是"Java"但引用地址不同。s4 = "Ja" + "va"是两个编译期常量拼接，编译器优化为"Java"并放入常量池，与s1指向同一对象。第一个输出s1 == s2：比较引用地址，s1和s2都指向常量池同一对象→true。第二个输出s1 == s3：s1指向常量池，s3指向堆上新对象，引用不同→false。第三个输出s1.equals(s3)：equals方法比较字符串的内容而非引用，内容都是"Java"→true。第四个输出s1 == s4：s4经编译器优化后指向常量池中同一个"Java"对象→true。第五个输出s1.length()：返回字符串长度4。'
        },
        {
            'score': '6',
            'code': '''public class WhileCopy {
    public static void main(String[] args) {
        String src = "HelloJavaWorld";
        StringBuffer result = new StringBuffer();
        int i = 0;
        char c = src.charAt(i);
        while (c != 'J') {
            result.append(c);
            c = src.charAt(++i);
        }
        System.out.println(result);
    }
}''',
            'answer': 'Hello',
            'analysis': "从main方法开始运行，运行到String src = \"HelloJavaWorld\"定义源字符串，看向StringBuffer result = new StringBuffer()创建空的可变字符串用于累积结果。int i = 0初始化索引，char c = src.charAt(i)获取src索引0处的字符'H'。进入while循环，循环条件是c != 'J'（当前字符不等于'J'）。在循环体中执行了result.append(c)将当前字符追加到result中，然后执行c = src.charAt(++i)（先自增i再获取下一个字符）。逐次追踪：第1轮c='H'（不是'J'），追加'H'，i变为1，c变为'e'；第2轮c='e'，追加'e'，i变为2，c变为'l'；第3轮c='l'，追加'l'，i变为3，c变为'l'；第4轮c='l'，追加'l'，i变为4，c变为'o'；第5轮c='o'，追加'o'，i变为5，c变为'J'；回到while条件判断c != 'J'为false，但不执行循环体（'J'本身不被追加到result中）。最后输出result的内容\"Hello\"——即'J'之前的所有字符。"
        },
        {
            'score': '6',
            'code': '''class Student {
    String name;
    int age;

    public Student(String name, int age) {
        this.name = name;
        this.age = age;
    }

    public void printInfo() {
        System.out.println(name + ", " + age + " years old");
    }
}

public class TestStudent {
    public static void main(String[] args) {
        Student s1 = new Student("张三", 20);
        Student s2 = new Student("李四", 22);
        s1.printInfo();
        s2.printInfo();
    }
}''',
            'answer': '''张三, 20 years old
李四, 22 years old''',
            'analysis': '从main方法开始运行，首先运行到Student s1 = new Student("张三", 20)，在这里面执行了new Student和构造方法Student(String name, int age)。在构造方法中this.name = name将参数字符串"张三"的引用赋给当前对象的name成员变量，this.age = age将20赋给age。因此s1指向的Student对象name为"张三"、age为20。接着运行Student s2 = new Student("李四", 22)，同理创建第二个对象s2，name为"李四"、age为22。然后执行s1.printInfo()，看向printInfo方法，在其中执行了System.out.println(name + ", " + age + " years old")，此处name是s1的name即"张三"，age是20，因此输出"张三, 20 years old"。s2.printInfo()同理，输出"李四, 22 years old"。'
        },
        {
            'score': '6',
            'code': '''public class DogTest {
    public static void main(String[] args) {
        Dog d1 = new Dog("旺财");
        Dog d2 = new Dog("小黑");
        d1.setName("大黄");
        System.out.println("Dog1: " + d1.getName());
        System.out.println("Dog2: " + d2.getName());
    }
}

class Dog {
    String name;
    Dog(String name) { this.name = name; }
    String getName() { return name; }
    void setName(String name) { this.name = name; }
}''',
            'answer': '''Dog1: 大黄
Dog2: 小黑''',
            'analysis': '从main方法开始运行，首先执行Dog d1 = new Dog("旺财")，调用Dog类的构造方法Dog(String name)，在其中执行了this.name = name将"旺财"赋给d1的name成员变量。接着Dog d2 = new Dog("小黑")同理创建d2，name为"小黑"。然后执行d1.setName("大黄")，看向setName方法，在这里面执行了this.name = name，其中this指向d1对象，参数name为"大黄"。关键点：this.name = name这行代码中，左边的this.name是成员变量，右边不带this的name是方法参数，两者虽然同名但是不同的变量。执行后d1的name成员变量从"旺财"被修改为"大黄"。d2没有调用setName，其name保持为"小黑"。第一个输出语句中d1.getName()返回"大黄"，输出"Dog1: 大黄"；第二个输出d2.getName()返回"小黑"，输出"Dog2: 小黑"。这里的重要知识点是this关键字的用法——在setName方法中，this.name指成员变量，而参数name遮蔽了成员变量名，this用于区分两者。'
        },
    ],
    'correction': [
        {
            'score': '10',
            'code': '''class Calculator {
    private int result;

    public void add(int a, int b) {
        result = a + b;
    }

    public int getResult() {
        return result;
    }
}

public class TestCalc {
    public static void main(String[] args) {
        Calculator c = new Calculator();
        c.add(10, 20);
        c.result = 100;  // 直接修改
        System.out.println(c.getResult());
    }
}''',
            'errors': [
                {'desc': '在TestCalc的main方法中直接访问c.result：result是Calculator类的private成员变量，不能在类外部直接访问，编译错误。',
                 'fix': '去掉c.result = 100这一行，或为Calculator添加public void setResult(int r)方法。'}
            ],
        },
        {
            'score': '10',
            'code': '''class Person {
    String name;
    public Person(String name) {
        this.name = name;
    }
}

class Student extends Person {
    int grade;
    // 编译错误：未定义构造方法
}

public class TestPerson {
    public static void main(String[] args) {
        Student s = new Student("小明");
        s.grade = 3;
        System.out.println(s.name);
    }
}''',
            'errors': [
                {'desc': 'Student类继承了Person，但Person只有带参构造方法Person(String name)。Student类没有显式定义构造方法，编译器会生成默认无参构造方法，其中隐式调用super()，但Person没有无参构造方法，导致编译错误。此外，new Student("小明")试图调用Student的带参构造方法，但Student未定义。',
                 'fix': '在Student类中添加构造方法：public Student(String name) { super(name); }'}
            ],
        },
    ],
    'design': [
        {
            'score': '20',
            'title': '温度转换工具类',
            'requirement': '''编写一个类 TemperatureConverter，要求：
- 包含两个静态方法：
  - public static double celsiusToFahrenheit(double celsius)：将摄氏温度转换为华氏温度（公式：F = C × 9/5 + 32）
  - public static double fahrenheitToCelsius(double fahrenheit)：将华氏温度转换为摄氏温度（公式：C = (F - 32) × 5/9）
- 编写 main 方法测试：输出0°C、100°C对应的华氏度，以及32°F、212°F对应的摄氏度。''',
            'reference_code': '''public class TemperatureConverter {
    public static double celsiusToFahrenheit(double celsius) {
        return celsius * 9.0 / 5.0 + 32;
    }

    public static double fahrenheitToCelsius(double fahrenheit) {
        return (fahrenheit - 32) * 5.0 / 9.0;
    }

    public static void main(String[] args) {
        System.out.println("0°C = " + celsiusToFahrenheit(0) + "°F");
        System.out.println("100°C = " + celsiusToFahrenheit(100) + "°F");
        System.out.println("32°F = " + fahrenheitToCelsius(32) + "°C");
        System.out.println("212°F = " + fahrenheitToCelsius(212) + "°C");
    }
}'''
        },
        {
            'score': '18',
            'title': '矩形类设计',
            'requirement': '''编写一个类 Rectangle，要求：
- 私有成员变量：double width（宽）、double height（高）
- 提供两个构造方法：
  - 无参构造方法：默认宽高均为 1.0
  - 有参构造方法：接收宽度和高度参数
- 提供方法：
  - double getArea()：返回面积
  - double getPerimeter()：返回周长
  - void printInfo()：打印矩形的宽、高、面积和周长
- 编写 main 方法，分别用无参和有参构造方法创建两个矩形对象并打印信息。''',
            'reference_code': '''public class Rectangle {
    private double width;
    private double height;

    public Rectangle() {
        this.width = 1.0;
        this.height = 1.0;
    }

    public Rectangle(double width, double height) {
        this.width = width;
        this.height = height;
    }

    public double getArea() {
        return width * height;
    }

    public double getPerimeter() {
        return 2 * (width + height);
    }

    public void printInfo() {
        System.out.println("Rectangle[width=" + width + ", height=" + height
            + ", area=" + getArea() + ", perimeter=" + getPerimeter() + "]");
    }

    public static void main(String[] args) {
        Rectangle r1 = new Rectangle();
        Rectangle r2 = new Rectangle(5.0, 3.0);
        r1.printInfo();
        r2.printInfo();
    }
}'''
        },
    ],
}


# ============================================================
# PAPER 7: 样卷7 (Intermediate Level)
# ============================================================
PAPER7 = {
    'title': 'Java程序设计 样卷7',
    'reading': [
        {
            'score': '6',
            'code': '''class Counter {
    private int serialNumber;
    public static int total = 0;

    public Counter() {
        total++;
        serialNumber = total;
    }

    public int getSerialNumber() {
        return serialNumber;
    }
}

public class TestCounter {
    public static void main(String[] args) {
        Counter c1 = new Counter();
        Counter c2 = new Counter();
        Counter c3 = new Counter();
        System.out.println("c1.serialNumber=" + c1.getSerialNumber());
        System.out.println("c2.serialNumber=" + c2.getSerialNumber());
        System.out.println("c3.serialNumber=" + c3.getSerialNumber());
        System.out.println("total=" + Counter.total);
    }
}''',
            'answer': '''c1.serialNumber=1
c2.serialNumber=2
c3.serialNumber=3
total=3''',
            'analysis': '从main方法开始运行，首先创建Counter[]数组。看向第一个new Counter()——执行Counter c1 = new Counter()，进入Counter构造方法：第一步执行total++，total是static变量，初始值为0，执行total++后total变为1；第二步serialNumber = total，将当前的total值1赋给c1的实例变量serialNumber。接着c2 = new Counter()，再次进入构造方法：执行total++，因为total是static变量（被所有实例共享，不从0重新开始），total从1变为2；serialNumber = total将2赋给c2。c3同理：total从2变为3，serialNumber赋值为3。三个输出语句分别打印c1.getSerialNumber()返回1、c2的返回2、c3的返回3。最后Counter.total直接通过类名访问静态变量total，输出3。这里的重要知识点：static变量属于类而非实例，只初始化一次，被所有实例共享；实例变量每个对象各有一份独立副本。如果total是实例变量（无static），则每个对象创建时都从0开始，结果将是1、1、1、但实际是1、2、3。'
        },
        {
            'score': '6',
            'code': '''class Shape {
    void draw() { System.out.println("Drawing Shape"); }
}

class Circle extends Shape {
    void draw() { System.out.println("Drawing Circle"); }
}

class Square extends Shape {
    void draw() { System.out.println("Drawing Square"); }
}

public class TestShape {
    static void drawShape(Shape s) {
        s.draw();
    }

    public static void main(String[] args) {
        Shape[] shapes = new Shape[3];
        shapes[0] = new Shape();
        shapes[1] = new Circle();
        shapes[2] = new Square();
        for (int i = 0; i < shapes.length; i++) {
            drawShape(shapes[i]);
        }
    }
}''',
            'answer': '''Drawing Shape
Drawing Circle
Drawing Square''',
            'analysis': '从main方法开始运行，首先创建Shape[] shapes = new Shape[3]数组（长度为3）。然后shapes[0] = new Shape()创建一个Shape对象，shapes[1] = new Circle()创建一个Circle对象（向上转型为Shape类型引用），shapes[2] = new Square()创建一个Square对象（同样向上转型）。进入for循环遍历数组，调用drawShape(shapes[i])。看向drawShape(Shape s)静态方法，它接收一个Shape类型的参数s，在方法体内执行了s.draw()。这里的关键是动态绑定（运行时多态）——虽然参数类型声明为Shape，但s实际指向的对象类型决定了调用哪个draw()。当s实际指向Shape对象时，调用Shape的draw()输出"Drawing Shape"；当s指向Circle对象时，调用Circle中重写的draw()输出"Drawing Circle"；当s指向Square对象时，调用Square中重写的draw()输出"Drawing Square"。编译时只知道s是Shape类型，但在运行时JVM根据对象的实际类型决定调用哪个方法版本——这就是多态的核心机制。'
        },
        {
            'score': '6',
            'code': '''class Parent {
    static { System.out.println("Parent static"); }
    { System.out.println("Parent instance"); }
    public Parent() { System.out.println("Parent constructor"); }
}

class Child extends Parent {
    static { System.out.println("Child static"); }
    { System.out.println("Child instance"); }
    public Child() { System.out.println("Child constructor"); }
}

public class TestInit {
    public static void main(String[] args) {
        System.out.println("---start---");
        Child c = new Child();
        System.out.println("---end---");
    }
}''',
            'answer': '''Parent static
Child static
---start---
Parent instance
Parent constructor
Child instance
Child constructor
---end---''',
            'analysis': '从main方法开始运行，首先执行System.out.println("---start---")输出---start---。但在此之前，由于main方法中要用到Child类，JVM在类加载阶段会先加载Parent类和Child类。加载Parent类时，执行static { System.out.println("Parent static"); }静态初始化块，输出"Parent static"。加载Child类时（Child继承Parent，但Parent已加载不会再执行静态块），执行Child的静态初始化块，输出"Child static"。所以最先看到的两行是静态块的输出（在---start---之前）。接下来执行Child c = new Child()创建子类对象，new的执行顺序是：第一步，进入Parent的实例初始化块{ System.out.println("Parent instance"); }，输出"Parent instance"；第二步，执行Parent的构造方法public Parent()，输出"Parent constructor"；第三步，进入Child的实例初始化块，输出"Child instance"；第四步，执行Child的构造方法，输出"Child constructor"。这个执行顺序体现了Java对象初始化的核心规则：先父后子，静态块（类加载时，仅一次）→实例块（每次new）→构造方法（每次new），且父类始终在子类之前。最后输出---end---。'
        },
        {
            'score': '6',
            'code': '''public class OverloadTest {
    public static int add(int a, int b) {
        return a + b;
    }

    public static double add(double a, double b) {
        return a + b;
    }

    public static String add(String a, String b) {
        return a + b;
    }

    public static void main(String[] args) {
        System.out.println(add(3, 5));
        System.out.println(add(2.5, 3.7));
        System.out.println(add("Hello", "World"));
    }
}''',
            'answer': '''8
6.2
HelloWorld''',
            'analysis': '从main方法开始运行，运行到第一个输出语句add(3, 5)，看向add方法的调用——编译器在编译时根据参数类型(3和5都是int)匹配到public static int add(int a, int b)版本，在其中执行了return a + b即3+5=8，输出8。第二个输出add(2.5, 3.7)，2.5和3.7都是double类型字面量，编译器匹配到public static double add(double a, double b)版本，返回2.5+3.7=6.2，输出6.2。第三个输出add("Hello", "World")，两个参数都是String类型，编译器匹配到public static String add(String a, String b)版本，在其中执行了return a + b，对于String类型来说+是字符串拼接运算符，结果"HelloWorld"，输出HelloWorld。这里的重要知识点是方法重载（Overloading）：同名方法根据参数类型和个数选择不同的版本，这是编译时多态——在编译阶段就已经确定了调用哪个方法版本，不是在运行时决定的。'
        },
        {
            'score': '6',
            'code': '''class Outer {
    private int x = 100;

    class Inner {
        void display() {
            System.out.println("x = " + x);
        }
    }

    void createInner() {
        Inner in = new Inner();
        in.display();
    }
}

public class TestInner {
    public static void main(String[] args) {
        Outer out = new Outer();
        out.createInner();
    }
}''',
            'answer': 'x = 100',
            'analysis': '从main方法开始运行，首先执行Outer out = new Outer()创建外部类对象，此时Outer的成员变量x被初始化为100。然后执行out.createInner()，看向createInner方法。在createInner方法中执行了Inner in = new Inner()创建内部类对象，然后调用in.display()。看向Inner类的display方法，在这里面执行了System.out.println("x = " + x)，其中x直接引用了外部类Outer的私有成员变量x。关键点：成员内部类（非静态内部类）可以直接访问外部类的所有成员（包括私有成员），因为内部类对象隐式持有一个指向外部类对象的引用（Outer.this）。Inner.display()中访问的x实际等价于Outer.this.x，所以输出"x = 100"。这种机制体现了内部类与外部类的紧密耦合关系——内部类是外部类的一个成员，可以直接访问外部类的状态。'
        },
        {
            'score': '6',
            'code': '''abstract class Animal {
    abstract void sound();
    public void sleep() {
        System.out.println("Animal sleeping");
    }
}

class Cat extends Animal {
    void sound() { System.out.println("Meow"); }
}

class Dog extends Animal {
    void sound() { System.out.println("Woof"); }
}

public class TestAnimal {
    public static void main(String[] args) {
        Animal a1 = new Cat();
        Animal a2 = new Dog();
        a1.sound();
        a1.sleep();
        a2.sound();
        a2.sleep();
    }
}''',
            'answer': '''Meow
Animal sleeping
Woof
Animal sleeping''',
            'analysis': '从main方法开始运行，首先执行Animal a1 = new Cat()——这里发生向上转型：Cat对象被赋值给Animal类型的引用a1。关键点：Animal是抽象类，不能直接实例化（new Animal()会编译错误），但可以作为引用类型指向子类对象。接着Animal a2 = new Dog()同理。然后a1.sound()：虽然a1的编译时类型是Animal，但运行时实际指向Cat对象，根据动态绑定机制，调用的是Cat类重写的sound()方法，输出"Meow"。a1.sleep()：sleep()是Animal中的具体方法，没有被Cat重写，因此调用的是Animal的sleep()，输出"Animal sleeping"。同样a2.sound()动态绑定到Dog的sound()，输出"Woof"；a2.sleep()调用继承的Animal.sleep()，输出"Animal sleeping"。这里体现了抽象类的设计模式：抽象方法定义"必须实现什么"（子类必须重写sound），具体方法提供"可以继承什么"（子类直接复用sleep）。'
        },
        {
            'score': '6',
            'code': '''public class BreakContinue {
    public static void main(String[] args) {
        for (int i = 1; i <= 5; i++) {
            if (i == 3) continue;
            if (i == 5) break;
            System.out.print(i + " ");
        }
        System.out.println();
        System.out.println("Done");
    }
}''',
            'answer': '''1 2 4
Done''',
            'analysis': '从main方法开始运行，进入for循环：i从1到5递增。循环体内有两个if判断。追踪每次循环：i=1时，if (i == 3)为false（不执行continue），if (i == 5)为false（不执行break），因此执行System.out.print(i + " ")输出"1 "。i=2时同理输出"2 "。i=3时，if (i == 3)为true，执行了continue语句——continue的作用是跳过本次循环剩余的所有代码，直接进入下一次循环的迭代（i++后判断条件）。因此对于i=3，不执行后面的if (i == 5)判断和print语句，直接跳到i=4。i=4时两个if都为false，输出"4 "。i=5时，if (i == 3)为false（不执行continue），但if (i == 5)为true，执行了break语句——break的作用是立即终止整个循环，不执行循环内剩余代码也不进行下一次迭代。因此不执行print语句，直接跳出for循环。循环结束后执行System.out.println()输出一个空行换行，最后输出"Done"。最终结果：打印"1 2 4 "然后换行输出"Done"。这里的关键是区分continue（跳过本次，继续下次循环）和break（直接跳出整个循环）的不同作用。'
        },
    ],
    'correction': [
        {
            'score': '10',
            'code': '''interface Drawable {
    void draw();
}

class Circle implements Drawable {
    // 错误：未实现接口方法
    void erase() {
        System.out.println("Erasing circle");
    }
}

public class TestDraw {
    public static void main(String[] args) {
        Drawable d = new Circle();
        d.draw();
    }
}''',
            'errors': [
                {'desc': 'Circle类声明实现了Drawable接口，但没有实现接口中声明的draw()方法。如果类不是抽象类，必须实现接口的所有方法，否则编译错误。',
                 'fix': '在Circle类中添加draw()方法的实现：public void draw() { System.out.println("Drawing circle"); }'}
            ],
        },
        {
            'score': '10',
            'code': '''class Parent {
    public static void info() {
        System.out.println("Parent info");
    }

    public void show() {
        System.out.println("Parent show");
    }
}

class Child extends Parent {
    public static void info() {
        System.out.println("Child info");
    }

    @Override
    public void show() {  // 正确重写
        System.out.println("Child show");
    }
}

public class TestOverride {
    public static void main(String[] args) {
        Parent p = new Child();
        p.info();    // 调用哪个？
        p.show();    // 调用哪个？
    }
}''',
            'errors': [
                {'desc': 'p.info()调用的是Parent.info()而非Child.info()。静态方法不能被重写（override），只能被隐藏（hide）。静态方法的调用在编译时根据引用类型决定，p是Parent类型，所以调用Parent.info()。这不是编译错误，但逻辑上容易误解。如果期望调用Child.info()，应使用Child.info()或修改引用类型为Child。',
                 'fix': '将静态方法改为实例方法（去掉static关键字），或使用Child.info()直接调用子类静态方法。'}
            ],
        },
    ],
    'design': [
        {
            'score': '20',
            'title': '银行账户体系',
            'requirement': '''编写程序，要求：
- 定义一个抽象类 Account，包含：
  - protected double balance（余额）
  - 构造方法初始化余额
  - 抽象方法 void deposit(double amount)（存款）
  - 抽象方法 void withdraw(double amount)（取款）
  - 具体方法 double getBalance() 返回余额
- 定义子类 SavingsAccount（储蓄账户）：
  - 存款时加 0.5% 利息
  - 取款时余额不足提示"余额不足"
- 定义子类 CreditAccount（信用账户）：
  - 存款无特殊处理
  - 取款时允许透支最多 1000 元
- 编写 main 方法，创建两种账户对象进行测试。''',
            'reference_code': '''abstract class Account {
    protected double balance;
    public Account(double balance) { this.balance = balance; }
    public abstract void deposit(double amount);
    public abstract void withdraw(double amount);
    public double getBalance() { return balance; }
}

class SavingsAccount extends Account {
    public SavingsAccount(double balance) { super(balance); }
    public void deposit(double amount) {
        balance += amount * 1.005;
        System.out.println("Savings deposit: " + amount + " -> balance: " + balance);
    }
    public void withdraw(double amount) {
        if (amount > balance) {
            System.out.println("余额不足");
        } else {
            balance -= amount;
            System.out.println("Savings withdraw: " + amount + " -> balance: " + balance);
        }
    }
}

class CreditAccount extends Account {
    public CreditAccount(double balance) { super(balance); }
    public void deposit(double amount) {
        balance += amount;
        System.out.println("Credit deposit: " + amount + " -> balance: " + balance);
    }
    public void withdraw(double amount) {
        if (amount > balance + 1000) {
            System.out.println("超过信用额度");
        } else {
            balance -= amount;
            System.out.println("Credit withdraw: " + amount + " -> balance: " + balance);
        }
    }
}

public class TestAccount {
    public static void main(String[] args) {
        SavingsAccount sa = new SavingsAccount(1000);
        sa.deposit(500);
        sa.withdraw(200);
        CreditAccount ca = new CreditAccount(500);
        ca.withdraw(1200);
        System.out.println("Final savings: " + sa.getBalance());
        System.out.println("Final credit: " + ca.getBalance());
    }
}'''
        },
        {
            'score': '18',
            'title': '自定义异常类',
            'requirement': '''编写程序，要求：
- 定义一个自定义异常类 ScoreException，继承 Exception：
  - 包含有参构造方法传递异常消息
- 定义一个类 ScoreValidator，包含静态方法：
  - public static void validate(int score) throws ScoreException
  - 如果 score < 0 或 score > 100，抛出 ScoreException，消息为"成绩不合法：X（成绩应在0-100之间）"
- 编写 main 方法：
  - 使用数组存放多个成绩 {95, 88, 105, 72, -5, 60}
  - 遍历数组并验证每个成绩，使用 try-catch 捕获异常并打印异常消息
  - 统计合法成绩的个数并输出。''',
            'reference_code': '''class ScoreException extends Exception {
    public ScoreException(String msg) { super(msg); }
}

class ScoreValidator {
    public static void validate(int score) throws ScoreException {
        if (score < 0 || score > 100) {
            throw new ScoreException("成绩不合法：" + score + "（成绩应在0-100之间）");
        }
    }
}

public class TestScore {
    public static void main(String[] args) {
        int[] scores = {95, 88, 105, 72, -5, 60};
        int validCount = 0;
        for (int s : scores) {
            try {
                ScoreValidator.validate(s);
                System.out.println("合法成绩：" + s);
                validCount++;
            } catch (ScoreException e) {
                System.out.println("异常：" + e.getMessage());
            }
        }
        System.out.println("合法成绩个数：" + validCount);
    }
}'''
        },
    ],
}


# ============================================================
# PAPER 8: 样卷8 (Upper-Intermediate)
# ============================================================
PAPER8 = {
    'title': 'Java程序设计 样卷8',
    'reading': [
        {
            'score': '6',
            'code': '''import java.util.*;

public class SetDemo {
    public static void main(String[] args) {
        Set<String> set = new HashSet<String>();
        set.add("apple");
        set.add("banana");
        set.add("apple");
        set.add("cherry");
        set.add("banana");
        System.out.println("size=" + set.size());
        System.out.println(set);
    }
}''',
            'answer': '''size=3
[banana, cherry, apple]（顺序可能不同）''',
            'analysis': '从main方法开始运行，首先创建HashSet<String>集合对象set，然后依次执行add操作。第一次set.add("apple")——HashSet内部通过hashCode和equals判断元素是否已存在，此时集合为空，直接添加成功，返回true。set.add("banana")同理添加成功。第二次set.add("apple")——关键点：此时集合中已经存在"apple"（String重写了hashCode和equals，内容相同的字符串hashCode相同且equals返回true），HashSet发现该元素已存在，因此不执行添加操作，add方法返回false（但程序没有检查返回值）。set.add("cherry")添加成功。第二次set.add("banana")同理，因为"banana"已存在，不添加。最终集合中只有3个不同元素：apple、banana、cherry。第一个输出size=3；第二个输出集合内容[banana, cherry, apple]或顺序可能不同——因为HashSet基于哈希表实现，不保证元素的迭代顺序与添加顺序一致。HashSet的核心特点：元素唯一、无序、允许null、基于hashCode实现快速查找。'
        },
        {
            'score': '6',
            'code': '''import java.util.*;

public class QueueDemo {
    public static void main(String[] args) {
        Queue<Integer> queue = new LinkedList<Integer>();
        for (int i = 3; i >= 0; i--) {
            queue.add(i);
        }
        System.out.println("Queue: " + queue);
        System.out.println("remove: " + queue.remove());
        System.out.println("remove: " + queue.remove());
        System.out.println("Queue: " + queue);
    }
}''',
            'answer': '''Queue: [3, 2, 1, 0]
remove: 3
remove: 2
Queue: [1, 0]''',
            'analysis': '从main方法开始运行，首先创建LinkedList作为Queue的实现，泛型为Integer。进入for循环，i从3递减到0，循环体中执行queue.add(i)将当前值添加到队列尾部。第1轮i=3：队列变为[3]（3在头部，也是尾部）；第2轮i=2：add将2添加到尾部，队列变为[3, 2]（3在头部，2在尾部）；第3轮i=1：队列变为[3, 2, 1]；第4轮i=0：队列变为[3, 2, 1, 0]。循环结束后第一个输出打印整个队列：[3, 2, 1, 0]。然后执行queue.remove()——remove方法从队列头部取出并删除元素，返回头部元素3，队列变为[2, 1, 0]，输出"remove: 3"。再次remove取出头部元素2，队列变为[1, 0]，输出"remove: 2"。最后打印队列显示剩余元素[1, 0]。关键点：队列遵循FIFO（先进先出）原则——add在尾部添加，remove从头部取出。LinkedList作为Queue的实现，完美支持这一特性。'
        },
        {
            'score': '6',
            'code': '''public class ExceptionDemo {
    public static void main(String[] args) {
        int[] arr = {10, 20, 30};
        try {
            System.out.println("Before loop");
            for (int i = 0; i <= arr.length; i++) {
                System.out.println("arr[" + i + "]=" + arr[i]);
            }
            System.out.println("After loop");
        } catch (ArrayIndexOutOfBoundsException e) {
            System.out.println("Exception caught: " + e.getMessage());
        }
        System.out.println("Program ends");
    }
}''',
            'answer': '''Before loop
arr[0]=10
arr[1]=20
arr[2]=30
Exception caught: Index 3 out of bounds for length 3（或类似消息）
Program ends''',
            'analysis': '从main方法开始运行，首先int[] arr = {10, 20, 30}创建长度为3的数组（索引0~2），然后进入try块。执行System.out.println("Before loop")输出第一行。进入for循环，循环条件是i <= arr.length即i <= 3。追踪循环：i=0时输出"arr[0]=10"；i=1时输出"arr[1]=20"；i=2时输出"arr[2]=30"；i=3时循环条件i<=3仍为true（因为数组长度是3），执行arr[3]时发生数组越界——arr的有效索引范围是0~2，arr[3]超出了数组边界，JVM抛出ArrayIndexOutOfBoundsException。此时try块内的后续代码（System.out.println("After loop")）不会被执行，控制流跳转到catch块。catch块捕获到该异常后执行System.out.println("Exception caught: " + e.getMessage())输出异常信息。这里注意catch处理后程序不会崩溃，继续执行catch块后面的代码System.out.println("Program ends")。关键点：循环条件i <= arr.length是错误写法（应使用i < arr.length），导致多循环了一次造成越界。try-catch保证了程序在异常发生后仍然能正常结束。'
        },
        {
            'score': '6',
            'code': '''class MyException extends Exception {
    public MyException(String msg) { super(msg); }
}

public class ThrowDemo {
    static void checkAge(int age) throws MyException {
        if (age < 0) {
            throw new MyException("Age cannot be negative: " + age);
        }
        System.out.println("Age is valid: " + age);
    }

    public static void main(String[] args) {
        try {
            checkAge(25);
            checkAge(-5);
        } catch (MyException e) {
            System.out.println("Error: " + e.getMessage());
        }
        System.out.println("Done");
    }
}''',
            'answer': '''Age is valid: 25
Error: Age cannot be negative: -5
Done''',
            'analysis': '从main方法开始运行，进入第一个try块调用checkAge(25)。看向checkAge方法，参数age=25，判断if (age < 0)为false（25不小于0），因此不执行throw语句，跳过if块后执行System.out.println("Age is valid: " + age)，输出"Age is valid: 25"，方法正常返回。第一个try块顺利结束（没有异常，跳过对应的catch）。进入第二个try块调用checkAge(-5)，参数age=-5，判断if (age < 0)为true，执行了throw new MyException("Age cannot be negative: " + age)——这里通过throw关键字主动抛出一个自定义异常对象，构造方法将"Age cannot be negative: -5"作为异常消息。抛出异常后，checkAge方法立即终止（后面的println不执行），异常向上传播到main方法的调用处。catch (MyException e)捕获该异常，执行System.out.println("Error: " + e.getMessage())输出"Error: Age cannot be negative: -5"。最后程序继续运行输出"Done"。关键点：throw用于手动抛出一个异常，方法通过throws声明告知调用者可能抛出什么异常，catch在调用处捕获处理。'
        },
        {
            'score': '6',
            'code': '''enum Season {
    SPRING("温暖"),
    SUMMER("炎热"),
    AUTUMN("凉爽"),
    WINTER("寒冷");

    private String desc;
    Season(String desc) { this.desc = desc; }
    public String getDesc() { return desc; }
}

public class EnumDemo {
    public static void main(String[] args) {
        Season s = Season.AUTUMN;
        System.out.println(s + " : " + s.getDesc());
        System.out.println("ordinal: " + s.ordinal());
        for (Season se : Season.values()) {
            System.out.print(se + " ");
        }
        System.out.println();
    }
}''',
            'answer': '''AUTUMN : 凉爽
ordinal: 2
SPRING SUMMER AUTUMN WINTER''',
            'analysis': '从main方法开始运行，首先执行Season s = Season.AUTUMN，将枚举常量AUTUMN赋值给引用s。第一个输出s + " : " + s.getDesc()：s.toString()返回"AUTUMN"（默认是枚举常量的名称），s.getDesc()调用AUTUMN实例的getDesc方法返回"凉爽"——因为AUTUMN("凉爽")在定义时通过构造方法将"凉爽"传给了desc成员变量。第二个输出s.ordinal()返回2——ordinal是枚举常量的序数值（从0开始），SPRING=0、SUMMER=1、AUTUMN=2、WINTER=3。第三个输出使用增强for循环遍历Season.values()，values()是编译器自动生成的静态方法，返回包含所有枚举常量的数组[SPRING, SUMMER, AUTUMN, WINTER]，依次输出每个枚举常量的名称（toString默认返回名称）。关键点：枚举enum本质上是一种特殊的类，可以有成员变量、构造方法（必须private）、方法；每个枚举常量都是该类的public static final实例。ordinal()返回定义时的顺序编号。'
        },
        {
            'score': '6',
            'code': '''import java.util.*;

public class ListDemo {
    public static void main(String[] args) {
        List<String> list = new ArrayList<String>();
        list.add("A");
        list.add("B");
        list.add("C");
        list.add(1, "X");
        list.remove("B");
        System.out.println("size=" + list.size());
        for (String s : list) {
            System.out.print(s + " ");
        }
        System.out.println();
    }
}''',
            'answer': '''size=3
A X C''',
            'analysis': '从main方法开始运行，首先创建ArrayList<String>对象list。依次执行操作：list.add("A")添加"A"到末尾→[A]；list.add("B")→[A, B]；list.add("C")→[A, B, C]。关键操作list.add(1, "X")：这是一个重载的add方法，第一个参数1是插入位置索引，第二个参数"X"是要插入的元素。执行后在索引1位置插入"X"，原索引1及之后的元素依次后移，列表变为[A, X, B, C]。接着list.remove("B")：根据内容删除第一个匹配的元素"B"，列表变为[A, X, C]。第一个输出size=3。第二个输出用增强for循环遍历列表依次打印每个元素，输出"A X C "。关键点：List的add有两个重载版本——add(E e)在末尾追加，add(int index, E e)在指定位置插入（索引从0开始）。remove(Object o)根据equals匹配删除第一个相等的元素。ArrayList允许重复元素并维护插入顺序。'
        },
        {
            'score': '6',
            'code': '''public class FinallyDemo {
    public static int test() {
        try {
            System.out.println("try");
            return 1;
        } finally {
            System.out.println("finally");
        }
    }

    public static void main(String[] args) {
        System.out.println("result=" + test());
        System.out.println("End");
    }
}''',
            'answer': '''try
finally
result=1
End''',
            'analysis': '从main方法开始运行，执行System.out.println("result=" + test())，这里需要先调用test()方法获取返回值。看向test()方法：进入try块，执行System.out.println("try")输出第一行"try"，然后执行return 1——但关键点来了：在return语句真正执行之前，JVM会先执行finally块中的代码。控制流进入finally块，执行System.out.println("finally")输出第二行"finally"。finally块执行完毕后，try块中的return 1才真正生效，test()方法返回1。回到main方法，拼接字符串输出"result=1"。最后输出"End"。所以执行顺序是try块的println→执行return但未返回→跳入finally块的println→finally执行完后return值生效→回到调用处。这里的重要知识点：finally块中的代码始终会被执行（即使try中有return），但finally不改变已经确定的返回值。如果finally中也有return，则会覆盖try中的return值。'
        },
    ],
    'correction': [
        {
            'score': '10',
            'code': '''import java.io.*;

public class FileWrite {
    public static void main(String[] args) {
        FileWriter fw = new FileWriter("output.txt");
        fw.write("Hello World");
        // 错误：未关闭流
        System.out.println("Write done");
    }
}''',
            'errors': [
                {'desc': '1. FileWriter构造方法和write方法都可能抛出IOException（已检查异常），但main方法既没有try-catch捕获，也没有throws声明，编译错误。\n2. FileWriter使用后未调用close()关闭流，可能导致数据未完全写入磁盘和资源泄漏。',
                 'fix': '使用try-catch-finally包围，在finally中调用fw.close()确保流被关闭；或使用try-with-resources自动关闭。'}
            ],
            'corrected_code': '''import java.io.*;

public class FileWrite {
    public static void main(String[] args) {
        FileWriter fw = null;
        try {
            fw = new FileWriter("output.txt");
            fw.write("Hello World");
            System.out.println("Write done");
        } catch (IOException e) {
            e.printStackTrace();
        } finally {
            try {
                if (fw != null) fw.close();
            } catch (IOException e) {
                e.printStackTrace();
            }
        }
    }
}'''
        },
        {
            'score': '10',
            'code': '''import java.util.*;

public class GenericError {
    public static void main(String[] args) {
        List list = new ArrayList();
        list.add("Hello");
        list.add(123);

        for (Object obj : list) {
            String s = (String) obj;  // 强制转换
            System.out.println(s.toUpperCase());
        }
    }
}''',
            'errors': [
                {'desc': '1. List使用了原始类型（raw type），没有使用泛型参数，编译器会发出警告。\n2. list中添加了String和Integer两种类型，遍历时强制转换为String，当遍历到Integer(123)时会抛出ClassCastException（Integer无法转为String）。',
                 'fix': '使用泛型List<String>确保类型安全：List<String> list = new ArrayList<String>(); 删除list.add(123)或将其改为字符串。"'}
            ],
        },
    ],
    'design': [
        {
            'score': '20',
            'title': '文件字符统计工具',
            'requirement': '''编写一个类 FileCharCounter，要求：
- 包含静态方法 public static Map<Character, Integer> countChars(String filePath)
  - 使用 FileReader 和 BufferedReader 读取文本文件
  - 统计文件中所有字符（不含空格和换行）出现的次数
  - 返回一个 TreeMap<Character, Integer>（按字符字典序排列）
- 包含静态方法 public static void printStats(String filePath)
  - 调用 countChars 方法
  - 打印每个字符及其出现次数
  - 打印总字符数
- 在 main 方法中调用 printStats 测试
- 使用 try-catch-finally 正确处理 IOException。''',
            'reference_code': '''import java.io.*;
import java.util.*;

public class FileCharCounter {
    public static Map<Character, Integer> countChars(String filePath) {
        Map<Character, Integer> map = new TreeMap<>();
        BufferedReader br = null;
        try {
            br = new BufferedReader(new FileReader(filePath));
            int c;
            while ((c = br.read()) != -1) {
                char ch = (char) c;
                if (ch != ' ' && ch != '\\n' && ch != '\\r') {
                    map.put(ch, map.getOrDefault(ch, 0) + 1);
                }
            }
        } catch (IOException e) {
            e.printStackTrace();
        } finally {
            try { if (br != null) br.close(); } catch (IOException e) {}
        }
        return map;
    }

    public static void printStats(String filePath) {
        Map<Character, Integer> map = countChars(filePath);
        int total = 0;
        for (Map.Entry<Character, Integer> entry : map.entrySet()) {
            System.out.println("'" + entry.getKey() + "': " + entry.getValue());
            total += entry.getValue();
        }
        System.out.println("Total characters: " + total);
    }

    public static void main(String[] args) {
        printStats("test.txt");
    }
}'''
        },
        {
            'score': '18',
            'title': '泛型栈实现',
            'requirement': '''编写一个泛型类 GenericStack<T>，要求：
- 使用 ArrayList<T> 作为内部存储
- 提供方法：
  - void push(T item)：将元素压入栈顶
  - T pop()：弹出并返回栈顶元素（栈空时返回null）
  - T peek()：查看栈顶元素但不弹出（栈空时返回null）
  - boolean isEmpty()：判断栈是否为空
  - int size()：返回栈中元素个数
- 编写 main 方法：
  - 创建 GenericStack<Integer>，压入10、20、30
  - 依次弹出并打印所有元素
  - 创建 GenericStack<String>，压入"A"、"B"、"C"
  - 查看栈顶元素并打印
  - 弹出所有元素并打印。''',
            'reference_code': '''import java.util.ArrayList;

public class GenericStack<T> {
    private ArrayList<T> list = new ArrayList<>();

    public void push(T item) {
        list.add(item);
    }

    public T pop() {
        if (list.isEmpty()) return null;
        return list.remove(list.size() - 1);
    }

    public T peek() {
        if (list.isEmpty()) return null;
        return list.get(list.size() - 1);
    }

    public boolean isEmpty() {
        return list.isEmpty();
    }

    public int size() {
        return list.size();
    }

    public static void main(String[] args) {
        GenericStack<Integer> intStack = new GenericStack<>();
        intStack.push(10);
        intStack.push(20);
        intStack.push(30);
        while (!intStack.isEmpty()) {
            System.out.println("Pop: " + intStack.pop());
        }

        GenericStack<String> strStack = new GenericStack<>();
        strStack.push("A");
        strStack.push("B");
        strStack.push("C");
        System.out.println("Peek: " + strStack.peek());
        while (!strStack.isEmpty()) {
            System.out.println("Pop: " + strStack.pop());
        }
    }
}'''
        },
    ],
}


# ============================================================
# PAPER: 提高卷4 (Advanced Level)
# ============================================================
ADVANCED4 = {
    'title': 'Java程序设计 提高卷4',
    'reading': [
        {
            'score': '6',
            'code': '''import java.io.*;

public class DataIOTest {
    public static void main(String[] args) throws IOException {
        DataOutputStream dos = new DataOutputStream(
            new FileOutputStream("data.bin"));
        dos.writeInt(100);
        dos.writeDouble(3.14);
        dos.writeUTF("Hello");
        dos.close();

        DataInputStream dis = new DataInputStream(
            new FileInputStream("data.bin"));
        System.out.println(dis.readInt());
        System.out.println(dis.readDouble());
        System.out.println(dis.readUTF());
        dis.close();
    }
}''',
            'answer': '''100
3.14
Hello''',
            'analysis': '从main方法开始运行，首先创建DataOutputStream对象包装FileOutputStream。执行dos.writeInt(100)将int值100以二进制格式写入data.bin文件（占4字节）；dos.writeDouble(3.14)将double值3.14以二进制格式写入（占8字节）；dos.writeUTF("Hello")将字符串"Hello"以UTF-8变长编码格式写入（先写长度信息再写内容）；dos.close()关闭输出流并刷新缓冲区。然后创建DataInputStream对象包装FileInputStream读取同一个文件data.bin。第一个System.out.println(dis.readInt())——readInt从文件中读取4个字节并解析为int值，返回100（对应之前写入的writeInt(100)）；第二个输出dis.readDouble()读取8个字节解析为double，返回3.14；第三个输出dis.readUTF()读取UTF编码的字符串，返回"Hello"。最后dis.close()关闭输入流。这里的关键点是DataOutputStream和DataInputStream必须严格配对使用，写入和读取的顺序、类型必须完全一致——如果顺序错误（比如先读double再读int），会得到错误的数据甚至异常，因为不同类型占用的字节数和编码方式不同。'
        },
        {
            'score': '6',
            'code': '''import java.util.*;

class Student implements Comparable<Student> {
    String name;
    int score;

    public Student(String name, int score) {
        this.name = name; this.score = score;
    }

    public int compareTo(Student o) {
        return this.score - o.score;
    }

    public String toString() {
        return name + "(" + score + ")";
    }
}

public class SortDemo {
    public static void main(String[] args) {
        List<Student> list = new ArrayList<>();
        list.add(new Student("Alice", 85));
        list.add(new Student("Bob", 92));
        list.add(new Student("Charlie", 78));
        Collections.sort(list);
        System.out.println(list);
    }
}''',
            'answer': '[Charlie(78), Alice(85), Bob(92)]',
            'analysis': '从main方法开始运行，首先创建ArrayList<Student>列表，依次添加三个Student对象：new Student("Alice", 85)、new Student("Bob", 92)、new Student("Charlie", 78)。列表当前顺序为[Alice(85), Bob(92), Charlie(78)]。然后执行Collections.sort(list)对列表排序。看向Student类，它实现了Comparable<Student>接口并重写了compareTo方法——return this.score - o.score，这意味着排序规则是当前对象的score减去另一个对象的score。Collections.sort内部使用compareTo方法比较元素：比较Alice(85)和Bob(92)，85-92=-7（负数，Alice排在Bob前面）；比较Bob(92)和Charlie(78)，92-78=14（正数，Bob排在Charlie后面）；以此类推进行排序。最终按score从小到大排列：[Charlie(78), Alice(85), Bob(92)]。最后输出列表，由于Student重写了toString方法返回"name(score)"格式，所以列表打印为[Charlie(78), Alice(85), Bob(92)]。关键点：compareTo返回负数表示this排在o前面，返回正数表示this排在o后面，返回0表示相等。'
        },
        {
            'score': '6',
            'code': '''import java.io.*;

public class CopyBytes {
    public static void main(String[] args) throws IOException {
        FileInputStream in = null;
        FileOutputStream out = null;
        try {
            in = new FileInputStream("source.txt");
            out = new FileOutputStream("target.txt");
            int c;
            int count = 0;
            while ((c = in.read()) != -1) {
                out.write(c);
                count++;
            }
            System.out.println("Copied " + count + " bytes");
        } finally {
            if (in != null) in.close();
            if (out != null) out.close();
        }
    }
}''',
            'answer': 'Copied <文件字节数> bytes（假设source.txt存在且有内容）',
            'analysis': '从main方法开始运行，首先声明FileInputStream和FileOutputStream引用并初始化为null（因为在try块中创建，需要在finally中访问）。进入try块，创建FileInputStream读取"source.txt"，创建FileOutputStream写入"target.txt"。int count = 0初始化计数器。进入while循环，条件是(c = in.read()) != -1——in.read()每次从源文件读取一个字节（返回0~255的int值），当读到文件末尾时返回-1。循环体中执行out.write(c)将读取到的字节c写入目标文件，然后count++计数。循环持续进行直到read()返回-1，此时while条件为false退出循环。循环结束后输出"Copied <count> bytes"。然后进入finally块（无论try块是否出现异常，finally都会执行），在这里面执行了if (in != null) in.close()和if (out != null) out.close()关闭流。使用if (in != null)判断是因为如果创建流时抛出异常，in可能仍然是null。关键点：使用try-finally确保即使发生异常流也能被正确关闭，避免资源泄漏；提前将流引用声明在try外部并赋null，是为了finally块中能访问到。'
        },
        {
            'score': '6',
            'code': '''class Animal {
    String name = "Animal";
    public String getName() { return name; }
    public void speak() { System.out.println("Animal speaks"); }
}

class Cat extends Animal {
    String name = "Cat";
    public String getName() { return name; }
    public void speak() { System.out.println("Meow"); }
}

public class FieldPoly {
    public static void main(String[] args) {
        Animal a = new Cat();
        System.out.println(a.name);
        System.out.println(a.getName());
        a.speak();
    }
}''',
            'answer': '''Animal
Cat
Meow''',
            'analysis': '从main方法开始运行，执行Animal a = new Cat()发生向上转型——Cat对象被赋值给Animal类型的引用a。此时a的编译时类型是Animal，运行时实际类型是Cat。第一个输出a.name：访问成员变量。关键点来了——Java中成员变量没有多态特性！变量的访问由编译时类型（引用类型）决定，而不是运行时类型。虽然a实际指向Cat对象，但因为a被声明为Animal类型，编译器直接访问Animal类中定义的name（值为"Animal"），而不是Cat类中的name（值为"Cat"）。所以输出"Animal"。第二个输出a.getName()：调用方法。方法调用有多态特性（动态绑定）——虽然a是Animal类型引用，但运行时JVM查看a实际指向的是Cat对象，因此调用Cat类中重写的getName()方法，返回Cat的name即"Cat"。第三个输出a.speak()：同理多态调用Cat的speak()，输出"Meow"。本程序的核心知识点是字段访问与方法调用的关键区别：字段访问是静态绑定的（编译时根据引用类型决定），方法调用是动态绑定的（运行时根据对象实际类型决定）。这就是为什么a.name输出"Animal"但a.getName()输出"Cat"的原因。'
        },
        {
            'score': '6',
            'code': '''import java.util.*;

public class MapDemo {
    public static void main(String[] args) {
        Map<String, Integer> map = new HashMap<>();
        map.put("one", 1);
        map.put("two", 2);
        map.put("three", 3);
        map.put("one", 10);

        System.out.println("Size: " + map.size());
        System.out.println("get one: " + map.get("one"));
        System.out.println("get four: " + map.get("four"));

        for (String key : map.keySet()) {
            System.out.println(key + " -> " + map.get(key));
        }
    }
}''',
            'answer': '''Size: 3
get one: 10
get four: null
（以下三行顺序可能不同）
one -> 10
two -> 2
three -> 3''',
            'analysis': '从main方法开始运行，创建HashMap<String, Integer>对象map。依次执行put操作：map.put("one", 1)将键"one"与值1关联→{one=1}；put("two", 2)→{one=1, two=2}；put("three", 3)→{one=1, two=2, three=3}。关键操作map.put("one", 10)：键"one"已经存在于map中（前面put过"one"→1），HashMap中键是唯一的——put相同的键不会添加新条目，而是用新值10覆盖旧值1，返回被覆盖的旧值1（但程序没有接收返回值）。此时map变为{one=10, two=2, three=3}。第一个输出size=3（虽然put了4次，但只有3个不同的键）。第二个输出get("one")返回10（被覆盖后的值）。第三个输出get("four")——键"four"不存在于map中，HashMap的get方法对不存在的键返回null（注意：这不会抛出异常，但后续如果对null做操作可能引发NullPointerException）。最后使用增强for循环遍历keySet()（返回所有键的集合），依次打印每个键和对应的值。由于HashMap不保证迭代顺序，输出顺序可能与put顺序不同（如可能先输出two->2，然后是three->3，最后one->10）。关键点：HashMap的put操作是覆盖而非追加，通过hashCode和equals判断键是否重复。'
        },
        {
            'score': '6',
            'code': '''class Outer {
    private static int sx = 10;
    private int ix = 20;

    static class StaticInner {
        void display() {
            System.out.println("sx=" + sx);
            // System.out.println("ix=" + ix);  // 编译错误
        }
    }

    class InstanceInner {
        void display() {
            System.out.println("sx=" + sx + ", ix=" + ix);
        }
    }
}

public class InnerDemo {
    public static void main(String[] args) {
        Outer.StaticInner si = new Outer.StaticInner();
        si.display();

        Outer o = new Outer();
        Outer.InstanceInner ii = o.new InstanceInner();
        ii.display();
    }
}''',
            'answer': '''sx=10
sx=10, ix=20''',
            'analysis': '从main方法开始运行，首先执行Outer.StaticInner si = new Outer.StaticInner()——这里创建静态内部类的对象不需要先创建外部类对象，直接用"外部类.静态内部类"的方式new，因为静态内部类属于外部类本身而非外部类的实例。然后si.display()，看向StaticInner的display方法，在里面执行了System.out.println("sx=" + sx)访问了Outer的静态成员sx（值为10），这是允许的，因为静态内部类与外部类的静态成员属于同一层级（都关联到类而非实例），输出"sx=10"。而被注释掉的System.out.println("ix=" + ix)是无法编译的——静态内部类不持有外部类实例的引用，因此不能直接访问外部类的实例成员ix（编译错误）。接着执行Outer o = new Outer()先创建外部类对象，然后Outer.InstanceInner ii = o.new InstanceInner()——实例内部类的创建必须通过外部类对象来new（语法：外部类对象.new 内部类()）。ii.display()，看向InstanceInner的display方法，执行System.out.println("sx=" + sx + ", ix=" + ix)，在这里可以同时访问静态成员sx=10和实例成员ix=20（因为o已创建），输出"sx=10, ix=20"。核心知识点：静态内部类（static class）不持有外部类引用，只能访问外部类静态成员；实例内部类持有外部类引用（Outer.this），可以访问外部类的所有成员。创建方式也不同。'
        },
        {
            'score': '6',
            'code': '''public class ExceptionChain {
    static void methodA() throws Exception {
        try {
            throw new ArithmeticException("Division by zero");
        } catch (ArithmeticException e) {
            throw new Exception("Error in methodA", e);
        }
    }

    public static void main(String[] args) {
        try {
            methodA();
        } catch (Exception e) {
            System.out.println("Message: " + e.getMessage());
            System.out.println("Cause: " + e.getCause());
        }
    }
}''',
            'answer': '''Message: Error in methodA
Cause: java.lang.ArithmeticException: Division by zero''',
            'analysis': '从main方法开始运行，进入try块调用methodA()。看向methodA方法：进入方法的try块，执行throw new ArithmeticException("Division by zero")手动抛出一个ArithmeticException异常对象（注意throw和throws的区别——throw是抛出异常的动作，throws是方法签名中的声明）。异常抛出后try块后续代码不执行，控制流跳转到catch块。catch (ArithmeticException e)捕获了这个异常，在catch块中执行了throw new Exception("Error in methodA", e)——这里又抛出另一个异常，但使用了Exception的双参构造方法Exception(String message, Throwable cause)，将捕获到的ArithmeticException对象e作为第二个参数（cause）传入，形成了异常链（chained exception）。这个新的Exception被抛出后向上传播到main方法。main方法的catch (Exception e)捕获到它，第一个输出e.getMessage()返回"Error in methodA"（构造时传入的第一个参数message）；第二个输出e.getCause()返回原始的ArithmeticException对象（toString显示"java.lang.ArithmeticException: Division by zero"）。关键点：异常链允许在高层包装低层异常的同时保留原始异常信息，getCause()可以追溯到根本原因。'
        },
    ],
    'correction': [
        {
            'score': '10',
            'code': '''import java.sql.*;

public class JDBCTest {
    public static void main(String[] args) {
        Connection conn = DriverManager.getConnection(
            "jdbc:mysql://localhost:3306/test", "root", "123456");
        Statement stmt = conn.createStatement();
        ResultSet rs = stmt.executeQuery("SELECT * FROM users");
        while (rs.next()) {
            System.out.println(rs.getString("name"));
        }
        // 错误：未关闭资源，也未处理异常
    }
}''',
            'errors': [
                {'desc': '1. JDBC操作需要处理SQLException（已检查异常），方法签名缺少throws声明或try-catch块。\n2. 需要加载驱动Class.forName("com.mysql.jdbc.Driver")（某些版本可省略但建议保留）。\n3. Connection、Statement、ResultSet使用后未关闭，应按先开后关的顺序在finally中关闭，否则造成资源泄漏。',
                 'fix': '添加try-catch-finally块，在finally中依次关闭ResultSet、Statement、Connection；或使用try-with-resources自动关闭资源。'}
            ],
            'corrected_code': '''import java.sql.*;

public class JDBCTest {
    public static void main(String[] args) {
        Connection conn = null;
        Statement stmt = null;
        ResultSet rs = null;
        try {
            Class.forName("com.mysql.jdbc.Driver");
            conn = DriverManager.getConnection(
                "jdbc:mysql://localhost:3306/test", "root", "123456");
            stmt = conn.createStatement();
            rs = stmt.executeQuery("SELECT * FROM users");
            while (rs.next()) {
                System.out.println(rs.getString("name"));
            }
        } catch (Exception e) {
            e.printStackTrace();
        } finally {
            try { if (rs != null) rs.close(); } catch (SQLException e) {}
            try { if (stmt != null) stmt.close(); } catch (SQLException e) {}
            try { if (conn != null) conn.close(); } catch (SQLException e) {}
        }
    }
}'''
        },
        {
            'score': '10',
            'code': '''import java.io.*;

class Person implements Serializable {
    String name;
    int age;
    // 错误：没有 serialVersionUID

    public Person(String name, int age) {
        this.name = name; this.age = age;
    }
}

public class SerialTest {
    public static void main(String[] args) throws Exception {
        Person p = new Person("Tom", 25);
        ObjectOutputStream oos = new ObjectOutputStream(
            new FileOutputStream("person.ser"));
        oos.writeObject(p);
        // 未调用 flush() 和 close()

        ObjectInputStream ois = new ObjectInputStream(
            new FileInputStream("person.ser"));
        Person p2 = (Person) ois.readObject();
        System.out.println(p2.name);
    }
}''',
            'errors': [
                {'desc': '1. ObjectOutputStream和ObjectInputStream使用后未关闭，可能导致数据未完全刷新和资源泄漏。应在finally中调用close()，或使用try-with-resources。\n2. writeObject后建议调用flush()确保缓冲数据写入磁盘。\n3. 类名SerialTest与文件名可能不一致（若文件名为SerialTest.java则正确，但需注意public类与文件名匹配）。',
                 'fix': '使用try-with-resources确保流自动关闭，或添加finally块手动关闭。添加flush()调用。'}
            ],
        },
    ],
    'design': [
        {
            'score': '20',
            'title': 'JDBC学生管理系统',
            'requirement': '''编写程序，使用JDBC连接MySQL数据库。要求：
- 数据库名为 School，表名为 Student，包含字段：
  - id INT PRIMARY KEY AUTO_INCREMENT
  - name VARCHAR(50)
  - age INT
  - score DOUBLE
- 定义方法：
  - void insertStudent(String name, int age, double score)：插入一条记录
  - void queryAllStudents()：查询所有记录并打印（id, name, age, score）
  - void queryHighScore(double minScore)：查询成绩大于等于minScore的学生并打印
- 在main方法中调用以上方法测试
- 使用try-catch-finally正确处理SQLException并关闭Connection、Statement、ResultSet
- 数据库连接：jdbc:mysql://127.0.0.1:3306/School，用户名root，密码123456。''',
            'reference_code': '''import java.sql.*;

public class StudentManager {
    private static final String URL = "jdbc:mysql://127.0.0.1:3306/School";
    private static final String USER = "root";
    private static final String PASS = "123456";

    public void insertStudent(String name, int age, double score) {
        Connection conn = null;
        PreparedStatement pstmt = null;
        try {
            Class.forName("com.mysql.jdbc.Driver");
            conn = DriverManager.getConnection(URL, USER, PASS);
            String sql = "INSERT INTO Student(name, age, score) VALUES(?, ?, ?)";
            pstmt = conn.prepareStatement(sql);
            pstmt.setString(1, name);
            pstmt.setInt(2, age);
            pstmt.setDouble(3, score);
            int rows = pstmt.executeUpdate();
            System.out.println("Inserted " + rows + " row(s)");
        } catch (Exception e) {
            e.printStackTrace();
        } finally {
            try { if (pstmt != null) pstmt.close(); } catch (SQLException e) {}
            try { if (conn != null) conn.close(); } catch (SQLException e) {}
        }
    }

    public void queryAllStudents() {
        Connection conn = null;
        Statement stmt = null;
        ResultSet rs = null;
        try {
            Class.forName("com.mysql.jdbc.Driver");
            conn = DriverManager.getConnection(URL, USER, PASS);
            stmt = conn.createStatement();
            rs = stmt.executeQuery("SELECT * FROM Student");
            System.out.println("ID\\tName\\tAge\\tScore");
            while (rs.next()) {
                System.out.println(rs.getInt("id") + "\\t"
                    + rs.getString("name") + "\\t"
                    + rs.getInt("age") + "\\t"
                    + rs.getDouble("score"));
            }
        } catch (Exception e) {
            e.printStackTrace();
        } finally {
            try { if (rs != null) rs.close(); } catch (SQLException e) {}
            try { if (stmt != null) stmt.close(); } catch (SQLException e) {}
            try { if (conn != null) conn.close(); } catch (SQLException e) {}
        }
    }

    public void queryHighScore(double minScore) {
        Connection conn = null;
        PreparedStatement pstmt = null;
        ResultSet rs = null;
        try {
            Class.forName("com.mysql.jdbc.Driver");
            conn = DriverManager.getConnection(URL, USER, PASS);
            String sql = "SELECT * FROM Student WHERE score >= ?";
            pstmt = conn.prepareStatement(sql);
            pstmt.setDouble(1, minScore);
            rs = pstmt.executeQuery();
            while (rs.next()) {
                System.out.println(rs.getString("name") + ": " + rs.getDouble("score"));
            }
        } catch (Exception e) {
            e.printStackTrace();
        } finally {
            try { if (rs != null) rs.close(); } catch (SQLException e) {}
            try { if (pstmt != null) pstmt.close(); } catch (SQLException e) {}
            try { if (conn != null) conn.close(); } catch (SQLException e) {}
        }
    }

    public static void main(String[] args) {
        StudentManager sm = new StudentManager();
        sm.insertStudent("张三", 20, 88.5);
        sm.insertStudent("李四", 21, 92.0);
        sm.queryAllStudents();
        sm.queryHighScore(90.0);
    }
}'''
        },
        {
            'score': '18',
            'title': '对象序列化工具',
            'requirement': '''编写程序，要求：
- 定义一个类 Book，实现 Serializable 接口：
  - 属性：String title、String author、double price、transient String password（不参与序列化）
  - 包含构造方法和 toString() 方法
- 定义一个工具类 SerialUtil，包含两个静态方法：
  - public static void saveBook(Book book, String filePath)：将Book对象序列化保存到文件
  - public static Book loadBook(String filePath)：从文件反序列化读取Book对象
- 在main方法中：
  - 创建Book对象，设置password为"abc123"
  - 序列化保存到"book.ser"
  - 从文件反序列化读取Book对象
  - 打印读取的对象，观察password字段的值
- 正确处理IOException和ClassNotFoundException。''',
            'reference_code': '''import java.io.*;

class Book implements Serializable {
    private static final long serialVersionUID = 1L;
    private String title;
    private String author;
    private double price;
    private transient String password;

    public Book(String title, String author, double price, String password) {
        this.title = title;
        this.author = author;
        this.price = price;
        this.password = password;
    }

    @Override
    public String toString() {
        return "Book[title=" + title + ", author=" + author
            + ", price=" + price + ", password=" + password + "]";
    }
}

class SerialUtil {
    public static void saveBook(Book book, String filePath) {
        ObjectOutputStream oos = null;
        try {
            oos = new ObjectOutputStream(new FileOutputStream(filePath));
            oos.writeObject(book);
            System.out.println("Book saved.");
        } catch (IOException e) {
            e.printStackTrace();
        } finally {
            try { if (oos != null) oos.close(); } catch (IOException e) {}
        }
    }

    public static Book loadBook(String filePath) {
        ObjectInputStream ois = null;
        try {
            ois = new ObjectInputStream(new FileInputStream(filePath));
            return (Book) ois.readObject();
        } catch (IOException | ClassNotFoundException e) {
            e.printStackTrace();
            return null;
        } finally {
            try { if (ois != null) ois.close(); } catch (IOException e) {}
        }
    }
}

public class TestBookSerial {
    public static void main(String[] args) {
        Book b1 = new Book("Java编程思想", "Bruce Eckel", 99.0, "abc123");
        System.out.println("Before: " + b1);
        SerialUtil.saveBook(b1, "book.ser");
        Book b2 = SerialUtil.loadBook("book.ser");
        System.out.println("After: " + b2);
    }
}'''
        },
    ],
}


# ============================================================
# Generate all files
# ============================================================
if __name__ == '__main__':
    # Paper 6
    create_paper(PAPER6, '样卷6.docx')
    create_answer(PAPER6, '样卷6答案.docx')

    # Paper 7
    create_paper(PAPER7, '样卷7.docx')
    create_answer(PAPER7, '样卷7答案.docx')

    # Paper 8
    create_paper(PAPER8, '样卷8.docx')
    create_answer(PAPER8, '样卷8答案.docx')

    # Advanced 4
    create_paper(ADVANCED4, '提高卷4.docx')
    create_answer(ADVANCED4, '提高卷4答案.docx')

    print('\nAll files generated successfully!')
